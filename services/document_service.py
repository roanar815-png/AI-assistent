"""
Сервис для работы с документами и шаблонами
"""
from docx import Document
from docx.shared import Pt
from docxtpl import DocxTemplate
from typing import Dict, List, Optional
import os
import shutil
from jinja2 import Template
from datetime import datetime
import uuid


class DocumentService:
    """Сервис для создания документов из шаблонов"""
    
    def __init__(self):
        """Инициализация сервиса"""
        self.templates_dir = "templates/documents"
        self.generated_dir = "generated_documents"
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.generated_dir, exist_ok=True)
    
    def fill_complaint_template(self, user_data: Dict) -> str:
        """
        Заполнить шаблон жалобы
        
        Args:
            user_data: Данные пользователя
        
        Returns:
            Путь к созданному документу
        """
        doc = Document()
        
        # Заголовок
        heading = doc.add_heading('ЖАЛОБА', 0)
        heading.alignment = 1  # Центрирование
        
        # Основной текст
        doc.add_paragraph(f"От: {user_data.get('full_name', 'Не указано')}")
        doc.add_paragraph(f"Организация: {user_data.get('organization', 'Не указано')}")
        doc.add_paragraph(f"Email: {user_data.get('email', 'Не указано')}")
        doc.add_paragraph(f"Телефон: {user_data.get('phone', 'Не указано')}")
        doc.add_paragraph("")
        
        doc.add_heading('Суть жалобы:', level=2)
        doc.add_paragraph(user_data.get('complaint_text', 'Текст жалобы не указан'))
        doc.add_paragraph("")
        
        doc.add_paragraph(f"Дата: {user_data.get('date', '')}")
        doc.add_paragraph(f"Подпись: _________________")
        
        # Сохранение
        filename = f"complaint_{user_data.get('full_name', 'unknown').replace(' ', '_')}.docx"
        filepath = f"generated_documents/{filename}"
        doc.save(filepath)
        
        return filepath
    
    def fill_protocol_template(self, meeting_data: Dict) -> str:
        """
        Заполнить шаблон протокола
        
        Args:
            meeting_data: Данные о собрании
        
        Returns:
            Путь к созданному документу
        """
        doc = Document()
        
        # Заголовок
        heading = doc.add_heading('ПРОТОКОЛ ЗАСЕДАНИЯ', 0)
        heading.alignment = 1
        
        # Детали
        doc.add_paragraph(f"Дата: {meeting_data.get('date', '')}")
        doc.add_paragraph(f"Время: {meeting_data.get('time', '')}")
        doc.add_paragraph(f"Место: {meeting_data.get('location', '')}")
        doc.add_paragraph("")
        
        doc.add_heading('Присутствовали:', level=2)
        participants = meeting_data.get('participants', [])
        for participant in participants:
            doc.add_paragraph(f"• {participant}", style='List Bullet')
        doc.add_paragraph("")
        
        doc.add_heading('Повестка дня:', level=2)
        doc.add_paragraph(meeting_data.get('agenda', 'Не указано'))
        doc.add_paragraph("")
        
        doc.add_heading('Решения:', level=2)
        doc.add_paragraph(meeting_data.get('decisions', 'Не указано'))
        doc.add_paragraph("")
        
        doc.add_paragraph(f"Председатель: {meeting_data.get('chairman', '')}")
        doc.add_paragraph(f"Секретарь: {meeting_data.get('secretary', '')}")
        
        # Сохранение
        filename = f"protocol_{meeting_data.get('date', 'unknown').replace(' ', '_')}.docx"
        filepath = f"generated_documents/{filename}"
        doc.save(filepath)
        
        return filepath
    
    def fill_contract_template(self, contract_data: Dict) -> str:
        """
        Заполнить шаблон договора
        
        Args:
            contract_data: Данные договора
        
        Returns:
            Путь к созданному документу
        """
        doc = Document()
        
        # Заголовок
        heading = doc.add_heading('ДОГОВОР О ЧЛЕНСТВЕ', 0)
        heading.alignment = 1
        
        # Номер и дата
        doc.add_paragraph(f"№ {contract_data.get('contract_number', '___')}")
        doc.add_paragraph(f"г. Москва, {contract_data.get('date', '')}")
        doc.add_paragraph("")
        
        # Стороны
        doc.add_heading('Стороны договора:', level=2)
        doc.add_paragraph(
            f"Организация «Опора России», в лице {contract_data.get('opora_representative', '[Представитель]')}, "
            f"действующего на основании Устава, именуемая в дальнейшем «Организация», с одной стороны,"
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            f"и {contract_data.get('member_name', '[ФИО]')}, "
            f"представляющий {contract_data.get('company_name', '[Название компании]')}, "
            f"ИНН {contract_data.get('inn', '[ИНН]')}, "
            f"именуемый в дальнейшем «Член», с другой стороны,"
        )
        doc.add_paragraph("")
        
        # Предмет договора
        doc.add_heading('1. Предмет договора', level=2)
        doc.add_paragraph(
            "Настоящий договор определяет условия членства в Организации «Опора России» "
            "и взаимные обязательства сторон."
        )
        doc.add_paragraph("")
        
        doc.add_heading('2. Обязанности сторон', level=2)
        doc.add_paragraph("2.1. Организация обязуется:")
        doc.add_paragraph("• Предоставлять информационную поддержку", style='List Bullet')
        doc.add_paragraph("• Организовывать мероприятия для членов", style='List Bullet')
        doc.add_paragraph("• Защищать интересы малого и среднего бизнеса", style='List Bullet')
        doc.add_paragraph("")
        doc.add_paragraph("2.2. Член обязуется:")
        doc.add_paragraph("• Оплачивать членские взносы", style='List Bullet')
        doc.add_paragraph("• Соблюдать Устав организации", style='List Bullet')
        doc.add_paragraph("")
        
        # Подписи
        doc.add_paragraph("")
        doc.add_paragraph("Организация: _________________")
        doc.add_paragraph("Член: _________________")
        
        # Сохранение
        filename = f"contract_{contract_data.get('member_name', 'unknown').replace(' ', '_')}.docx"
        filepath = f"generated_documents/{filename}"
        doc.save(filepath)
        
        return filepath
    
    def create_custom_document(self, template_text: str, 
                              data: Dict) -> str:
        """
        Создать документ из текстового шаблона с подстановкой данных
        
        Args:
            template_text: Текст шаблона с плейсхолдерами {{variable}}
            data: Данные для подстановки
        
        Returns:
            Путь к созданному документу
        """
        template = Template(template_text)
        filled_text = template.render(**data)
        
        doc = Document()
        for paragraph in filled_text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        filename = f"custom_document_{data.get('user_id', 'unknown')}.docx"
        filepath = f"generated_documents/{filename}"
        doc.save(filepath)
        
        return filepath
    
    def upload_template(self, file_path: str, template_name: str, 
                       description: str = "") -> Dict:
        """
        Загрузить шаблон документа
        
        Args:
            file_path: Путь к загружаемому файлу
            template_name: Название шаблона
            description: Описание шаблона
        
        Returns:
            Информация о загруженном шаблоне
        """
        try:
            # Генерируем уникальное имя файла
            file_ext = os.path.splitext(file_path)[1]
            template_id = str(uuid.uuid4())
            template_filename = f"{template_id}{file_ext}"
            template_path = os.path.join(self.templates_dir, template_filename)
            
            # Копируем файл
            shutil.copy2(file_path, template_path)
            
            # Создаем метаданные
            metadata = {
                "template_id": template_id,
                "name": template_name,
                "description": description,
                "filename": template_filename,
                "original_filename": os.path.basename(file_path),
                "upload_date": datetime.now().isoformat(),
                "file_type": file_ext[1:] if file_ext else "unknown"
            }
            
            # Сохраняем метаданные
            metadata_path = os.path.join(self.templates_dir, f"{template_id}.json")
            import json
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
            
            return {
                "status": "success",
                "template_id": template_id,
                "message": f"Шаблон '{template_name}' успешно загружен"
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка загрузки шаблона: {str(e)}"
            }
    
    def get_templates_list(self) -> List[Dict]:
        """
        Получить список всех шаблонов
        
        Returns:
            Список шаблонов с метаданными
        """
        templates = []
        try:
            for filename in os.listdir(self.templates_dir):
                if filename.endswith('.json'):
                    metadata_path = os.path.join(self.templates_dir, filename)
                    import json
                    with open(metadata_path, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)
                    templates.append(metadata)
        except Exception as e:
            print(f"Ошибка получения списка шаблонов: {e}")
        
        return templates
    
    def fill_uploaded_template(self, template_id: str, user_data: Dict, 
                              conversation_data: Dict = None) -> str:
        """
        Заполнить загруженный шаблон данными
        
        Args:
            template_id: ID шаблона
            user_data: Данные пользователя
            conversation_data: Данные из разговора
        
        Returns:
            Путь к созданному документу
        """
        try:
            # Загружаем метаданные
            metadata_path = os.path.join(self.templates_dir, f"{template_id}.json")
            import json
            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
            
            template_path = os.path.join(self.templates_dir, metadata['filename'])
            
            # Подготавливаем данные для подстановки
            fill_data = {
                **user_data,
                "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "user_id": user_data.get("user_id", "unknown")
            }
            
            if conversation_data:
                fill_data.update({
                    "message": conversation_data.get("message", ""),
                    "response": conversation_data.get("response", "")
                })
            
            # Обрабатываем в зависимости от типа файла
            if metadata['file_type'] == 'docx':
                return self._fill_docx_template(template_path, fill_data, metadata)
            elif metadata['file_type'] == 'txt':
                return self._fill_text_template(template_path, fill_data, metadata)
            else:
                raise ValueError(f"Неподдерживаемый тип файла: {metadata['file_type']}")
                
        except Exception as e:
            print(f"Ошибка заполнения шаблона: {e}")
            raise
    
    def _fill_docx_template(self, template_path: str, data: Dict, 
                           metadata: Dict) -> str:
        """
        Заполнить DOCX шаблон с ПОЛНЫМ сохранением структуры (docxtpl)
        Сохраняет ВСЕ: таблицы, рамки, линии, shapes, text boxes, форматирование
        """
        try:
            print(f"\n[FILL DOCX] ЗАПОЛНЕНИЕ DOCX ШАБЛОНА (docxtpl):", flush=True)
            print(f"   Шаблон: {metadata.get('name', 'Unknown')}", flush=True)
            print(f"   Путь: {template_path}", flush=True)
            print(f"   Данных для заполнения: {len([k for k, v in data.items() if v])}", flush=True)
            
            # Показываем какие данные есть
            filled_fields = {k: v for k, v in data.items() if v and v != ""}
            print(f"   Заполненные поля ({len(filled_fields)}):")
            for key, value in sorted(filled_fields.items()):
                if key != 'user_id':
                    print(f"      • {key}: {str(value)[:60]}")
            
            # Открываем шаблон через docxtpl (сохраняет ВСЮ структуру!)
            print(f"\n[LOADING] Загружаем шаблон через DocxTemplate...", flush=True)
            doc = DocxTemplate(template_path)
            
            # Подготавливаем контекст для рендеринга
            # Преобразуем None в пустые строки
            context = {}
            for key, value in data.items():
                if value is None or value == "":
                    context[key] = ""
                else:
                    context[key] = str(value)
            
            # ДОБАВЛЯЕМ МАППИНГ ПОЛЕЙ для совместимости с разными шаблонами
            if 'full_name' in context and context['full_name']:
                context['fio'] = context['full_name']  # Для шаблонов, использующих {{fio}}
                context['name'] = context['full_name']  # Для шаблонов, использующих {{name}}
                context['фио'] = context['full_name']  # Для шаблонов на русском
                context['имя'] = context['full_name']  # Альтернативное поле
            
            if 'organization' in context and context['organization']:
                context['org'] = context['organization']  # Сокращенное название
                context['организация'] = context['organization']  # На русском
            
            print(f"\n[RENDER] Заполняем шаблон данными...", flush=True)
            print(f"   Передаем {len(context)} полей в контекст", flush=True)
            
            # МАГИЯ! docxtpl автоматически заполняет ВСЕ элементы документа:
            # - Параграфы, таблицы, заголовки, колонтитулы
            # - Text boxes, shapes, рамки, линии
            # - Сохраняет ВСЁ форматирование, границы, стили
            doc.render(context)
            
            print(f"   [OK] Шаблон заполнен! Все структуры сохранены.", flush=True)
            
            # Сохраняем результат (заменяем кириллические символы для совместимости с веб-интерфейсом)
            safe_name = metadata['name'].replace(' ', '_').replace('Документ', 'Document').replace('Россия', 'Russia')
            output_filename = f"{safe_name}_{data.get('user_id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.generated_dir, output_filename)
            
            print(f"\n[SAVE] Сохраняем документ: {output_filename}", flush=True)
            doc.save(output_path)
            
            # Возвращаем путь с прямыми слешами для веб-интерфейса
            web_path = output_path.replace('\\', '/')
            return web_path
            
        except Exception as e:
            print(f"Ошибка заполнения DOCX шаблона: {e}")
            raise
    
    def _fill_text_template(self, template_path: str, data: Dict, 
                           metadata: Dict) -> str:
        """
        Заполнить текстовый шаблон и создать DOCX
        """
        try:
            # Читаем шаблон
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
            
            # Заполняем плейсхолдеры
            template = Template(template_content)
            filled_content = template.render(**data)
            
            # Создаем DOCX документ
            doc = Document()
            for line in filled_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Сохраняем (заменяем кириллические символы для совместимости с веб-интерфейсом)
            safe_name = metadata['name'].replace(' ', '_').replace('Документ', 'Document').replace('Россия', 'Russia')
            output_filename = f"{safe_name}_{data.get('user_id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.generated_dir, output_filename)
            doc.save(output_path)
            
            # Возвращаем путь с прямыми слешами для веб-интерфейса
            web_path = output_path.replace('\\', '/')
            return web_path
            
        except Exception as e:
            print(f"Ошибка заполнения текстового шаблона: {e}")
            raise
    
    def delete_template(self, template_id: str) -> Dict:
        """
        Удалить шаблон
        
        Args:
            template_id: ID шаблона
        
        Returns:
            Результат операции
        """
        try:
            # Загружаем метаданные
            metadata_path = os.path.join(self.templates_dir, f"{template_id}.json")
            import json
            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
            
            # Удаляем файлы
            template_path = os.path.join(self.templates_dir, metadata['filename'])
            if os.path.exists(template_path):
                os.remove(template_path)
            
            if os.path.exists(metadata_path):
                os.remove(metadata_path)
            
            return {
                "status": "success",
                "message": f"Шаблон '{metadata['name']}' удален"
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка удаления шаблона: {str(e)}"
            }


document_service = DocumentService()

