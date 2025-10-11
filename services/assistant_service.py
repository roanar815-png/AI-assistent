"""
Основной сервис ассистента для обработки диалогов
"""
from typing import Dict, List, Optional
from datetime import datetime
from models.schemas import ChatMessage, ChatResponse, UserData
from integrations import openai_service, google_sheets_service
from services.document_service import document_service
from logger_config import get_logger, log_success, log_error, log_warning

# Инициализация логгера
logger = get_logger(__name__)

# Глобальное хранилище сессий автозаполнения
_autofill_sessions = {}

class AssistantService:
    """Сервис для работы ассистента"""
    
    def __init__(self):
        """Инициализация сервиса"""
        self.conversations = {}  # Хранение истории диалогов в памяти
        self.autofill_sessions = _autofill_sessions  # Используем глобальное хранилище
    
    def process_message(self, user_id: str, message: str) -> ChatResponse:
        """
        Обработать сообщение пользователя
        
        Args:
            user_id: ID пользователя
            message: Текст сообщения
        
        Returns:
            Ответ ассистента
        """
        logger.info("=" * 80)
        logger.info("🤖 ОБРАБОТКА НОВОГО СООБЩЕНИЯ")
        logger.info(f"   User ID: {user_id}")
        logger.info(f"   Message: {message[:100]}{'...' if len(message) > 100 else ''}")
        logger.info("=" * 80)
        
        # Получить или создать историю диалога
        if user_id not in self.conversations:
            self.conversations[user_id] = []
            logger.debug(f"Создана новая сессия для пользователя: {user_id}")
        else:
            logger.debug(f"Продолжение сессии (история: {len(self.conversations[user_id])} сообщений)")
        
        conversation_history = self.conversations[user_id]
        
        # Проверка на специальные команды
        try:
            logger.debug("Отправка запроса к OpenAI/DeepSeek API...")
            if ("анализ мсп" in message.lower() or "прогноз мсп" in message.lower() or 
                "прогнозирование рынка" in message.lower() or "прогноз рынка" in message.lower()):
                logger.info("📊 Обнаружен запрос на анализ МСП")
                response_text = openai_service.analyze_sme_trends(message)
                action = "analysis"
            else:
                # Обычный диалог + попытка классифицировать намерение
                response_text = openai_service.chat(message, conversation_history)
                action = "chat"
            
            print(f"[SUCCESS] Получен ответ от AI (длина: {len(response_text)} символов)")
            print(f"   Первые 100 символов: {response_text[:100]}...")
            
            if not response_text or len(response_text.strip()) == 0:
                print(f"[WARNING] AI вернул ПУСТОЙ ответ!")
                response_text = "Извините, не удалось получить ответ от AI. Проверьте настройки API."
                
        except Exception as e:
            print(f"[ERROR] КРИТИЧЕСКАЯ ОШИБКА OpenAI/DeepSeek API: {e}")
            import traceback
            traceback.print_exc()
            response_text = f"Извините, произошла ошибка при обработке запроса: {str(e)}"
            action = "error"
        
        # Обновить историю
        conversation_history.append({"role": "user", "content": message})
        conversation_history.append({"role": "assistant", "content": response_text})
        
        # Сохранить в Google Sheets
        google_sheets_service.save_chat_history(user_id, message, response_text)
        
        # Попытка извлечь данные пользователя и бизнес-намерения
        self._extract_and_save_user_data(user_id, conversation_history)
        self._detect_intent_and_persist(user_id, conversation_history)
        
        # Проверяем, является ли сообщение жалобой
        self._check_and_save_complaint(user_id, message, conversation_history)
        
        # Проверяем, нужно ли создать документ из шаблона
        print(f"\n{'='*60}")
        print(f"ПРОВЕРКА СОЗДАНИЯ ДОКУМЕНТА")
        print(f"User ID: {user_id}")
        print(f"Message: {message[:100]}...")
        print(f"Response: {response_text[:100]}...")
        print(f"{'='*60}\n")
        
        try:
            document_suggestion = self._check_document_creation(user_id, message, response_text, conversation_history)
        except Exception as e:
            print(f"[ERROR] КРИТИЧЕСКАЯ ОШИБКА в _check_document_creation: {e}")
            import traceback
            traceback.print_exc()
            document_suggestion = None
        
        print(f"\n{'='*60}")
        print(f"[RESULT] document_suggestion:")
        if document_suggestion:
            print(f"   [+] document_suggestion создан")
            print(f"   - suggested: {document_suggestion.get('suggested')}")
            print(f"   - created_document exists: {document_suggestion.get('created_document') is not None}")
            if document_suggestion.get('created_document'):
                print(f"   - created_document.status: {document_suggestion.get('created_document', {}).get('status')}")
                print(f"   - created_document.filepath: {document_suggestion.get('created_document', {}).get('filepath')}")
                print(f"   - created_document.download_url: {document_suggestion.get('created_document', {}).get('download_url')}")
            print(f"   - needs_data: {document_suggestion.get('needs_data')}")
            print(f"   - message exists: {document_suggestion.get('message') is not None}")
        else:
            print(f"   [-] document_suggestion = None (не создан)")
        print(f"{'='*60}\n")
        
        # Если документ был создан автоматически, НЕ добавляем в текстовый ответ
        # Вместо этого передаем document_suggestion в ChatResponse
        if document_suggestion and document_suggestion.get("created_document"):
            created_doc = document_suggestion["created_document"]
            print(f"[PROCESSING] created_document:")
            print(f"   Status: {created_doc.get('status')}")
            if created_doc.get("status") == "success":
                print(f"   [OK] Статус success, документ будет отображен через document_suggestion")
                print(f"   Download URL: {created_doc.get('download_url')}")
                # НЕ добавляем в response_text - интерфейс обработает document_suggestion
            else:
                print(f"   [ERROR] Статус не success: {created_doc.get('status')}")
                # В случае ошибки добавляем информацию в текстовый ответ
                response_text += f"\n\n❌ Ошибка создания документа: {created_doc.get('message', 'Неизвестная ошибка')}"
        
        print(f"\n{'='*80}")
        print(f"[SENDING] ОТВЕТ КЛИЕНТУ:")
        print(f"   Response length: {len(response_text)}")
        print(f"   Action: {action}")
        print(f"   Document suggestion: {document_suggestion is not None}")
        print(f"   Response preview: {response_text[:150]}...")
        print(f"{'='*80}\n")
        
        return ChatResponse(
            response=response_text,
            action=action,
            document_suggestion=document_suggestion
        )
    
    def _extract_and_save_user_data(self, user_id: str, 
                                   conversation: List[Dict]):
        """
        Извлечь и сохранить данные пользователя из диалога
        
        Args:
            user_id: ID пользователя
            conversation: История диалога
        """
        # Формируем текст диалога
        conversation_text = "\n".join(
            [f"{msg['role']}: {msg['content']}" for msg in conversation[-10:]]
        )
        
        # Извлекаем данные
        user_info = openai_service.extract_user_info(conversation_text)
        
        # Если есть новые данные, сохраняем
        if user_info and any(user_info.values()):
            user_info['user_id'] = user_id
            google_sheets_service.save_user_data(user_info)

    def _detect_intent_and_persist(self, user_id: str, conversation: List[Dict]):
        """
        Определить намерение (заявка/документ/обратная связь) и сохранить данные
        """
        conversation_text = "\n".join(
            [f"{msg['role']}: {msg['content']}" for msg in conversation[-12:]]
        )
        intent_result = openai_service.detect_intent_and_extract(conversation_text) or {}

        intent = intent_result.get("intent")
        if not intent or intent == "none":
            return

        if intent == "application":
            application = intent_result.get("application") or {}
            if any(application.values()):
                # Дополняем user_id, сохраняем как заявку
                payload = {
                    "user_id": user_id,
                    "full_name": application.get("full_name") or "",
                    "email": application.get("email") or "",
                    "phone": application.get("phone") or "",
                    "organization": application.get("organization") or "",
                    "inn": application.get("inn") or "",
                    "business_type": application.get("business_type") or "",
                    "comment": application.get("comment") or ""
                }
                try:
                    google_sheets_service.save_application(payload)
                except Exception:
                    pass

        elif intent == "document":
            doc = intent_result.get("document") or {}
            template_type = doc.get("template_type")
            user_data = doc.get("user_data") or {}
            if template_type in {"complaint", "protocol", "contract"}:
                # Документ создается через document_service API, здесь только сохраняем пользователя
                try:
                    user_data = dict(user_data)
                    user_data["user_id"] = user_id
                    google_sheets_service.save_user_data(user_data)
                except Exception:
                    pass

        elif intent == "feedback":
            fb = intent_result.get("feedback") or {}
            if fb.get("message"):
                try:
                    google_sheets_service.save_feedback({
                        "user_id": user_id,
                        "message": fb.get("message"),
                        "category": fb.get("category") or "question"
                    })
                except Exception:
                    pass
    
    def _check_document_creation(self, user_id: str, message: str, 
                                 response: str, conversation: List[Dict]) -> Dict:
        """
        Проверить, нужно ли предложить создание документа из шаблона
        (РАСШИРЕННАЯ ВЕРСИЯ с анализом полноты данных и оценкой уверенности)
        
        Args:
            user_id: ID пользователя
            message: Сообщение пользователя
            response: Ответ ассистента
            conversation: История диалога
        
        Returns:
            Предложение о создания или None
        """
        print(f"\n{'='*80}")
        print(f"[CHECK DOC CREATION] _check_document_creation() вызвана")
        print(f"   User ID: {user_id}")
        print(f"   Message preview: {message[:100]}...")
        print(f"   Response preview: {response[:100]}...")
        print(f"{'='*80}\n")
        
        try:
            # Получаем список доступных шаблонов
            templates = document_service.get_templates_list()
            print(f"Найдено шаблонов: {len(templates)}")
            if templates:
                for t in templates:
                    print(f"   - {t.get('name', 'Без названия')} (ID: {t.get('template_id', 'N/A')[:8]}...)")
            
            if not templates:
                print(f"Нет доступных шаблонов, выход")
                return None
            
            # 🚨 АВАРИЙНАЯ ЛОГИКА: Прямое обнаружение запросов на создание документов
            message_lower = message.lower()
            emergency_trigger = (
                # Запрос "Заполни документ Россия" с данными
                (("заполни" in message_lower or "заполните" in message_lower) and
                "документ" in message_lower and
                "россия" in message_lower and
                (":" in message or "=" in message)) or
                
                # Запрос на вступление в Опору России с данными организации
                (("помогите подать заявку" in message_lower or "подать заявку" in message_lower) and
                ("вступление" in message_lower or "опору" in message_lower) and
                ("ооо" in message_lower or "общество" in message_lower or "инн" in message_lower or "огрн" in message_lower)) or
                
                # Запрос на вступление для физических лиц
                (("помогите подать заявление" in message_lower or "подать заявление" in message_lower) and
                ("вступление" in message_lower or "опору" in message_lower)) or
                
                # Запрос с данными ФИО (даже без двоеточий)
                (("помогите подать заявление" in message_lower or "подать заявление" in message_lower) and
                ("вступление" in message_lower or "опору" in message_lower) and
                ("зовут" in message_lower or "фамилия" in message_lower or "имя" in message_lower)) or
                
                # Запрос на создание жалобы с данными
                ("жалоб" in message_lower and
                ("фио" in message_lower or "телефон" in message_lower or "email" in message_lower or "адресат" in message_lower or "суть" in message_lower or "требования" in message_lower or "контактные данные" in message_lower))
            )
            
            if emergency_trigger:
                print(f"\n{'='*60}")
                print(f"[EMERGENCY MODE] ACTIVATED!")
                print(f"   Обнаружен запрос на создание документа с данными")
                print(f"   ОБХОД всех проверок AI - создаем документ ПРИНУДИТЕЛЬНО")
                print(f"{'='*60}\n")
                
                # Определяем тип запроса и выбираем подходящий шаблон
                is_organization_request = (
                    ("ооо" in message_lower or "общество" in message_lower or "инн" in message_lower or "огрн" in message_lower) or
                    ("помогите подать заявку" in message_lower and "вступление" in message_lower) or
                    ("подать заявку на вступление" in message_lower) or
                    ("помогите подать заявление" in message_lower and "вступление" in message_lower) or
                    ("подать заявление на вступление" in message_lower)
                )
                
                is_individual_request = (
                    ("помогите подать заявление" in message_lower and "вступление" in message_lower and 
                     ("зовут" in message_lower or "фамилия" in message_lower or "имя" in message_lower)) or
                    ("подать заявление на вступление" in message_lower and 
                     ("зовут" in message_lower or "фамилия" in message_lower or "имя" in message_lower))
                )
                
                # Специальная проверка для запроса "документ россия"
                is_russia_document_request = (
                    "документ" in message_lower and "россия" in message_lower
                )
                
                # Проверка для запроса на создание жалобы
                is_complaint_request = (
                    "жалоб" in message_lower and
                    ("фио" in message_lower or "телефон" in message_lower or "email" in message_lower or "адресат" in message_lower or "суть" in message_lower or "требования" in message_lower or "контактные данные" in message_lower)
                )
                
                # Находим подходящий шаблон
                selected_template = None
                
                # Приоритет 1: Если запрошена жалоба - создаем жалобу без шаблона
                if is_complaint_request:
                    print(f"[OK] Запрос на создание жалобы обнаружен")
                    # Создаем жалобу напрямую через document_service
                    try:
                        # Извлекаем данные для жалобы
                        import re
                        complaint_data = {"user_id": user_id}
                        
                        # Извлекаем ФИО
                        fio_match = re.search(r'ФИО[:\s]+([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)', message, re.IGNORECASE)
                        if fio_match:
                            complaint_data['full_name'] = fio_match.group(1)
                        
                        # Извлекаем телефон
                        phone_match = re.search(r'Телефон[:\s]+(\+?[0-9\s\(\)\-]+)', message, re.IGNORECASE)
                        if phone_match:
                            complaint_data['phone'] = phone_match.group(1)
                        
                        # Извлекаем email
                        email_match = re.search(r'E-mail[:\s]+([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', message, re.IGNORECASE)
                        if email_match:
                            complaint_data['email'] = email_match.group(1)
                        
                        # Извлекаем адресата
                        addressee_match = re.search(r'Адресат жалобы[:\s]+([^:]+?)(?=\s+Суть жалобы)', message, re.IGNORECASE | re.DOTALL)
                        if addressee_match:
                            complaint_data['addressee'] = addressee_match.group(1).strip()
                        
                        # Извлекаем суть жалобы
                        complaint_match = re.search(r'Суть жалобы[:\s]+([^:]+?)(?=\s+Требования)', message, re.IGNORECASE | re.DOTALL)
                        if complaint_match:
                            complaint_data['complaint_text'] = complaint_match.group(1).strip()
                        
                        # Извлекаем требования
                        requirements_match = re.search(r'Требования[:\s]+([^:]+?)(?=\s+Приложения)', message, re.IGNORECASE | re.DOTALL)
                        if requirements_match:
                            complaint_data['requirements'] = requirements_match.group(1).strip()
                        
                        # Создаем жалобу
                        filepath = document_service.fill_complaint_template(complaint_data)
                        
                        # Сохраняем в Google Sheets
                        google_sheets_service.save_complaint({
                            "complaint_id": f"COMP_{user_id}_{int(datetime.now().timestamp())}",
                            "user_id": user_id,
                            "full_name": complaint_data.get('full_name', ''),
                            "email": complaint_data.get('email', ''),
                            "phone": complaint_data.get('phone', ''),
                            "organization": complaint_data.get('organization', ''),
                            "complaint_text": complaint_data.get('complaint_text', ''),
                            "category": "Документированная жалоба",
                            "priority": "Средний"
                        })
                        
                        # Возвращаем результат
                        return {
                            "suggested": True,
                            "created_document": {
                                "status": "success",
                                "filepath": filepath,
                                "download_url": f"/api/documents/download?file={filepath}",
                                "template_name": "Жалоба",
                                "message": f" Жалоба успешно создана!"
                            },
                            "needs_data": False,
                            "message": f"**Жалоба готова!**\n\n**Документ:** Жалоба\n**Файл:** {filepath.split('/')[-1]}\n\n**Документ создан и готов к скачиванию.**"
                        }
                        
                    except Exception as e:
                        print(f"[ERROR] Ошибка создания жалобы: {e}")
                        import traceback
                        traceback.print_exc()
                
                # Приоритет 2: Если явно запрошен "документ россия" - ищем именно его
                elif is_russia_document_request:
                    for template in templates:
                        template_name_lower = template.get('name', '').lower()
                        if "документ" in template_name_lower and "россия" in template_name_lower:
                            selected_template = template
                            print(f"[OK] Найден шаблон 'Документ Россия': {template['name']}")
                            break
                
                # Приоритет 3: Для организаций ищем шаблон с "вступление" (заявление или анкета)
                elif is_organization_request:
                    for template in templates:
                        template_name_lower = template.get('name', '').lower()
                        if "вступление" in template_name_lower:
                            selected_template = template
                            print(f"[OK] Найден шаблон для организации: {template['name']}")
                            break
                
                # Приоритет 4: Для физических лиц ищем "Документ Россия"
                elif is_individual_request:
                    for template in templates:
                        template_name_lower = template.get('name', '').lower()
                        if "документ" in template_name_lower and "россия" in template_name_lower:
                            selected_template = template
                            print(f"[OK] Найден шаблон для физического лица: {template['name']}")
                            break
                
                # Приоритет 5: Fallback - ищем любой подходящий шаблон
                if not selected_template:
                    for template in templates:
                        template_name_lower = template.get('name', '').lower()
                        if "россия" in template_name_lower:
                            selected_template = template
                            print(f"[OK] Найден fallback шаблон: {template['name']}")
                            break
                
                if selected_template:
                    # Извлекаем ВСЕ данные напрямую регулярными выражениями
                    import re
                    user_info = {"user_id": user_id}
                    
                    # УЛУЧШЕННЫЕ паттерны - захватывают ВСЁ до следующего ИЗВЕСТНОГО поля
                    field_patterns = {
                        # Базовые поля - короткие
                        'last_name': r'(?:Фамилия|фамилия)[\s:=]+([А-ЯЁ][а-яё]+)',
                        'first_name': r'(?:Имя|имя)[\s:=]+([А-ЯЁа-яё]+)',
                        'middle_name': r'(?:Отчество|отчество)[\s:=]+([А-ЯЁ][а-яё]+)',
                        
                        # Альтернативные паттерны для фразы "Меня зовут Лев фамилия Балакин отчество Михайлович"
                        'first_name_alt': r'(?:зовут|имя)\s+([А-ЯЁа-яё]+)(?=\s+(?:фамилия|отчество|$))',
                        'last_name_alt': r'(?:фамилия)\s+([А-ЯЁ][а-яё]+)(?=\s+(?:отчество|$))',
                        'middle_name_alt': r'(?:отчество)\s+([А-ЯЁ][а-яё]+)',
                        
                        # Паттерн для фразы "мое ФИО Балакин Лев Михайлович"
                        'fio_direct': r'(?:мое\s+ФИО|ФИО|фио)[\s:=]+([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)',
                        'inn': r'(?:ИНН|инн)[\s:=]+(\d{10,12})',
                        'phone': r'(?:Телефон|телефон)[\s:=]+(\+?[0-9\s\(\)\-]+)',
                        'email': r'(?:E-mail|email|E-mail|почта)[\s:=]+([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                        
                        # Поля для организаций
                        'organization': r'(?:Полное наименование юридического лица|наименование)[\s:=]+([^:]+?)(?=\s+(?:ОГРН|ИНН|Юридический адрес))',
                        'ogrn': r'(?:ОГРН|огрн)[\s:=]+(\d{13,15})',
                        'legal_address': r'(?:Юридический адрес|юридический адрес)[\s:=]+([^:]+?)(?=\s+(?:Фактический адрес|Основной вид деятельности))',
                        'actual_address': r'(?:Фактический адрес|фактический адрес)[\s:=]+([^:]+?)(?=\s+(?:Основной вид деятельности|Контактные данные))',
                        'okved': r'(?:Основной вид деятельности|ОКВЭД)[\s:=]+([^:]+?)(?=\s+(?:Контактные данные|ФИО руководителя))',
                        'director_name': r'(?:ФИО руководителя|руководитель)[\s:=]+([^:]+?)(?=\s+(?:Должность|Телефон))',
                        'position': r'(?:Должность|должность)[\s:=]+([^:]+?)(?=\s+(?:Телефон|Электронная почта))',
                        'employees_count': r'(?:Количество сотрудников|сотрудников)[\s:=]+([^:]+?)(?=\s+(?:Годовой оборот|Регион деятельности))',
                        'annual_turnover': r'(?:Годовой оборот|оборот)[\s:=]+([^:]+?)(?=\s+(?:Регион деятельности|$))',
                        'activity_region': r'(?:Регион деятельности|регион)[\s:=]+([^:]+?)(?=\s*$)',
                        
                        # Длинные поля - захватываем ВСЁ до следующего известного ключевого слова
                        'birth_date': r'(?:Число[,\s]*месяц[,\s]*год\s+рождения|дата\s+рождения)[\s:=]+(.+?)(?=\s+(?:Адрес|Телефон|E-mail|Паспорт|Образование|Место\s+работы)|$)',
                        'region': r'(?:Регион|регион)[\s:=]+([А-ЯЁа-яё\s\-]+?)(?=\s+(?:Город|город))',
                        'city': r'(?:Город|город)[\s:=]+([А-ЯЁа-яё\-\s]+?)(?=\s+(?:Улица|улица))',
                        'street': r'(?:Улица|улица)[\s:=]+([А-ЯЁа-яё\s\-]+?)(?=\s+(?:Дом|дом))',
                        'house': r'(?:Дом|дом)[\s:=]+(\d+[а-яА-Яa-zA-Z]*)',
                        'apartment': r'(?:Кв\.?|квартира)[\s:=]+(\d+)',
                        
                        # Очень длинные поля - захватываем до следующего поля
                        'passport': r'(?:Паспортные\s+данные|паспорт)[\s:=]+(.+?)(?=\s+(?:Образование|образование):)',
                        'education': r'(?:Образование|образование)[\s:=]+(.+?)(?=\s+(?:Место\s+работы|место\s+работы),)',
                        'work_info': r'(?:Место\s+работы[,\s]*должность|работа)[\s:=]+(.+?)(?=\s+(?:Сфера\s+бизнеса|сфера\s+бизнеса))',
                        'activity_sphere': r'(?:Сфера\s+бизнеса[^:]*:|сфера[^:]*:)[\s:=]+(.+?)(?=\s+(?:Опыт\s+предпринимательской|опыт\s+предпринимательской))',
                        'business_experience': r'(?:Опыт\s+предпринимательской\s+деятельности|опыт\s+предпринимательской\s+деятельности)[\s:=]+(.+?)(?=\s+(?:Опыт\s+общественной|опыт\s+общественной))',
                        'public_activity_experience': r'(?:Опыт\s+общественной\s+деятельности|опыт\s+общественной\s+деятельности)[\s:=]+(.+?)(?=\s+(?:Эксперт\s+в\s+отрасли|эксперт\s+в\s+отрасли):)',
                        'expertise_area': r'(?:Эксперт\s+в\s+отрасли|эксперт\s+в\s+отрасли)[\s:=]+(.+?)(?=\s+(?:Выборные\s+должности|выборные\s+должности):)',
                        'elected_position': r'(?:Выборные\s+должности|выборные\s+должности)[\s:=]+(.+?)(?=\s+(?:Дополнения|дополнения):)',
                        'additional_info': r'(?:Дополнения|дополнения)[\s:=]+(.+?)(?=\s+(?:Дата\s+заполнения|дата\s+заполнения):)',
                    }
                    
                    for field_name, pattern in field_patterns.items():
                        # DOTALL позволяет .+? захватывать переносы строк
                        match = re.search(pattern, message, re.IGNORECASE | re.DOTALL)
                        if match:
                            value = match.group(1).strip()
                            # Убираем лишние точки и запятые в конце
                            value = value.rstrip('.,;:')
                            # Убираем лишние пробелы и переносы
                            value = ' '.join(value.split())
                            user_info[field_name] = value
                            print(f"   [OK] {field_name}: {value[:80]}...", flush=True)
                    
                    # Объединяем ФИО (основные и альтернативные паттерны)
                    # Сначала проверяем прямое извлечение ФИО
                    if user_info.get('fio_direct'):
                        user_info['fio'] = user_info['fio_direct']
                        user_info['full_name'] = user_info['fio_direct']
                        print(f"   [OK] fio (прямое извлечение): {user_info['fio']}")
                    else:
                        # Если нет прямого ФИО, собираем из частей
                        first_name = user_info.get('first_name') or user_info.get('first_name_alt')
                        last_name = user_info.get('last_name') or user_info.get('last_name_alt')
                        middle_name = user_info.get('middle_name') or user_info.get('middle_name_alt')
                        
                        if last_name or first_name:
                            parts = []
                            if last_name: parts.append(last_name)
                            if first_name: parts.append(first_name)
                            if middle_name: parts.append(middle_name)
                            user_info['full_name'] = ' '.join(parts)
                            user_info['fio'] = user_info['full_name']  # Добавляем поле fio для шаблонов
                            print(f"   [OK] full_name (объединено): {user_info['full_name']}")
                            print(f"   [OK] fio (для шаблонов): {user_info['fio']}")
                        
                        # Сохраняем в основных полях для совместимости
                        if not user_info.get('first_name') and first_name:
                            user_info['first_name'] = first_name
                        if not user_info.get('last_name') and last_name:
                            user_info['last_name'] = last_name
                        if not user_info.get('middle_name') and middle_name:
                            user_info['middle_name'] = middle_name
                    
                    # Объединяем адрес
                    if user_info.get('region') or user_info.get('city'):
                        address_parts = []
                        if user_info.get('region'): address_parts.append(user_info['region'])
                        if user_info.get('city'): address_parts.append(f"г. {user_info['city']}")
                        if user_info.get('street'): address_parts.append(f"ул. {user_info['street']}")
                        if user_info.get('house'):
                            house_str = f"д. {user_info['house']}"
                            if user_info.get('apartment'):
                                house_str += f", кв. {user_info['apartment']}"
                            address_parts.append(house_str)
                        user_info['address'] = ', '.join(address_parts)
                        print(f"   [OK] address (объединен): {user_info['address'][:70]}...")
                    
                    # СОЗДАЕМ ДОКУМЕНТ НЕМЕДЛЕННО
                    print(f"\n[CREATING] ДОКУМЕНТ В АВАРИЙНОМ РЕЖИМЕ...")
                    try:
                        created_document = self.create_document_from_template(
                            user_id=user_id,
                            template_id=selected_template['template_id'],
                            user_data=user_info,
                            conversation_data={
                                "message": message,
                                "response": response
                            },
                            send_email=True
                        )
                        
                        if created_document and created_document.get("status") == "success":
                            print(f"[OK] АВАРИЙНЫЙ РЕЖИМ: Документ создан успешно!")
                            print(f"   Filepath: {created_document.get('filepath')}")
                            print(f"   Download URL: {created_document.get('download_url')}")
                            
                            # ✅ СОХРАНЯЕМ В GOOGLE SHEETS
                            print(f"\n[GOOGLE SHEETS] Сохраняем документ в таблицу...", flush=True)
                            try:
                                google_sheets_service.save_document({
                                    "user_id": user_id,
                                    "full_name": user_info.get("full_name", ""),
                                    "email": user_info.get("email", ""),
                                    "document_type": "анкета",  # Для документа Россия
                                    "template_name": selected_template['name'],
                                    "filepath": created_document.get("filepath", ""),
                                    "download_url": created_document.get("download_url", ""),
                                    "completeness_score": 100,  # В аварийном режиме считаем что данные полные
                                    "confidence_score": 100,
                                    "data_quality": "excellent"
                                })
                                print(f"   [OK] Документ сохранен в Google Sheets!", flush=True)
                            except Exception as e:
                                print(f"   [WARNING] Ошибка сохранения в Google Sheets (не критично): {e}", flush=True)
                            
                            # Возвращаем результат
                            return {
                                "suggested": True,
                                "created_document": created_document,
                                "needs_data": False,
                                "message": f"""**Документ успешно создан!**

**Шаблон:** {selected_template['name']}
**Файл:** {created_document.get('filepath', '').split('/')[-1]}

**Ваш документ готов к скачиванию.**"""
                            }
                        else:
                            print(f"[ERROR] АВАРИЙНЫЙ РЕЖИМ: Не удалось создать документ")
                            print(f"   Результат: {created_document}")
                    except Exception as e:
                        print(f"[ERROR] ОШИБКА в аварийном режиме: {e}")
                        import traceback
                        traceback.print_exc()
                else:
                    print(f"[ERROR] АВАРИЙНЫЙ РЕЖИМ: Подходящий шаблон не найден")
            
            # Формируем текст для анализа (последние 5 сообщений)
            conversation_text = "\n".join(
                [f"{msg['role']}: {msg['content']}" for msg in conversation[-5:]]
            )
            
            # СТРОГИЙ СПИСОК: только явные запросы на ДЕЙСТВИЕ с документами
            document_keywords = [
                "создайте документ", "создать документ", "создай документ",
                "заполните анкету", "заполнить анкету", "заполни анкету",
                "оформите заявление", "оформить заявление",
                "подготовьте договор", "подготовить договор",
                "нужен документ", "нужна анкета", "нужно заявление",
                "хочу вступить", "подать заявку", "вступление в",
                "помогите создать", "помоги создать",
                "сделайте документ", "сделать документ",
                # Дополнительные ключевые слова для вступления в Опору России
                "помогите подать заявку", "помоги подать заявку",
                "подать заявку на вступление", "заявка на вступление",
                "вступление в опору", "вступить в опору",
                "помогите вступить", "помоги вступить",
                "заявка в опору россии", "вступление в опору россии",
                # Ключевые слова для физических лиц
                "помогите подать заявление", "помоги подать заявление",
                "подать заявление на вступление", "заявление на вступление"
            ]
            
            # УЛУЧШЕНИЕ: Проверяем ключевые слова в ПОСЛЕДНИХ 3 сообщениях, а не только в текущем
            recent_messages = " ".join([msg['content'].lower() for msg in conversation[-3:]])
            print(f"Анализ последних сообщений: '{recent_messages[:200]}...'")
            
            # Проверяем каждое ключевое слово в истории диалога
            found_keywords = [keyword for keyword in document_keywords if keyword in recent_messages]
            print(f"Найденные ключевые слова: {found_keywords}")
            
            # УЛУЧШЕНИЕ: Проверяем, не генерирует ли бот сам текст документа
            response_lower = response.lower()
            is_generating_document_text = False
            
            # Признаки генерации документа в ответе бота
            document_generation_indicators = [
                # Заголовки документов
                ("заявление" in response_lower and ("прошу принять" in response_lower or "от:" in response_lower)),
                ("анкета" in response_lower and ("фамилия:" in response_lower or "имя:" in response_lower)),
                ("договор" in response_lower and ("стороны:" in response_lower or "предмет договора" in response_lower)),
                
                # Множество полей в формате "Поле: Значение" или "Поле = Значение"
                (response_lower.count("фамилия:") > 0 and response_lower.count("имя:") > 0),
                (response_lower.count("телефон:") > 0 and response_lower.count("email:") > 0 and response_lower.count("адрес:") > 0),
                
                # Структура анкеты
                ("личные данные" in response_lower and "контактная информация" in response_lower),
            ]
            
            is_generating_document_text = any(document_generation_indicators)
            
            if is_generating_document_text:
                print("[WARNING] ОБНАРУЖЕНО: Бот генерирует текст документа вместо создания файла!")
            
            if is_generating_document_text:
                print("Бот генерирует текст документа вместо создания файла!")
                # Принудительно активируем создание документа
                found_keywords.append("документ (автообнаружение)")
            
            # КРИТИЧНО: Документы создаются ТОЛЬКО при ЯВНОМ намерении
            # Проверяем ТОЛЬКО сообщения пользователя, а не ответы бота
            user_messages_only = " ".join([msg['content'].lower() for msg in conversation[-3:] if msg['role'] == 'user'])
            print(f"Проверка ТОЛЬКО пользовательских сообщений: '{user_messages_only[:150]}...'")
            
            # СТРОГАЯ проверка: нужны ЯВНЫЕ ключевые слова создания документа
            document_action_keywords = [
                "создать", "создайте", "создай", 
                "заполнить", "заполните", "заполни",  # ДОБАВЛЕНО: заполни
                "оформить", "оформите", "оформи",
                "подготовить", "подготовьте", "подготовь",
                "сделать", "сделайте", "сделай",
                "нужен документ", "нужна анкета", "нужно заявление",
                "хочу вступить", "подать заявку", "вступление"
            ]
            
            has_document_action = any(kw in user_messages_only for kw in document_action_keywords)
            
            # НОВОЕ: Проверяем, есть ли в сообщении название шаблона + структурированные данные
            has_template_name = False
            detected_template_by_name = None
            
            # УЛУЧШЕНИЕ: Более гибкий поиск шаблона
            for template in templates:
                template_name = template.get('name', '')
                template_name_lower = template_name.lower()
                template_desc = template.get('description', '').lower()
                
                # Проверяем каждое слово из названия шаблона
                template_words = [w for w in template_name_lower.split() if len(w) > 3]
                
                # Подсчет совпадений
                matches = sum(1 for word in template_words if word in user_messages_only)
                
                # СПЕЦИАЛЬНАЯ ЛОГИКА для "Документ Россия"
                # Если в сообщении есть "документ" + "россия" + "анкета" + структурированные данные
                is_document_russia_request = (
                    "документ" in user_messages_only and 
                    "россия" in user_messages_only and
                    "россия" in template_name_lower
                )
                
                # Общая проверка: хотя бы 2 ключевых слова или специальный случай
                if matches >= 2 or (matches >= 1 and len(template_words) <= 2) or is_document_russia_request:
                    has_template_name = True
                    detected_template_by_name = template
                    print(f"[OK] Обнаружено название шаблона в сообщении: '{template.get('name')}'")
                    print(f"   Совпадений: {matches}, Специальная логика: {is_document_russia_request}")
                    break
            
            # Проверяем формат структурированных данных (Поле = Значение ИЛИ Поле: Значение)
            has_structured_data_equals = "=" in user_messages_only and len(user_messages_only.split("=")) >= 3
            has_structured_data_colon = ":" in user_messages_only and len(user_messages_only.split(":")) >= 3
            has_structured_data = has_structured_data_equals or has_structured_data_colon
            
            # Если есть название шаблона + структурированные данные = это запрос на заполнение!
            if has_template_name and has_structured_data:
                print(f"[DETECTED] Обнаружен запрос на заполнение: название шаблона + данные")
                print(f"   has_template_name: {has_template_name}")
                print(f"   has_structured_data: {has_structured_data} (equals: {has_structured_data_equals}, colon: {has_structured_data_colon})")
                print(f"   detected_template: {detected_template_by_name.get('name') if detected_template_by_name else 'None'}")
                has_document_action = True
            
            print(f"Явное действие с документом: {has_document_action}")
            
            # Если НЕТ явного действия - НЕ создаем документ
            if not has_document_action and not is_generating_document_text:
                print(f"Нет явного запроса на создание документа, выход")
                return None
            
            # Проверяем что это НЕ аналитический запрос
            analytical_keywords = [
                "прогноз", "прогнозирование", "анализ", "тренд", "статистика", 
                "аналитика", "динамика", "оценка", "обзор", "исследование",
                "покажите", "расскажите", "объясните", "как работает", "рынок"
            ]
            is_analytical_request = any(kw in user_messages_only for kw in analytical_keywords)
            
            print(f"[CHECK] Проверка типа запроса:")
            print(f"   is_analytical_request: {is_analytical_request}")
            print(f"   has_document_action: {has_document_action}")
            
            if is_analytical_request and not has_document_action:
                print(f"   [ERROR] Это аналитический/информационный запрос, документ НЕ нужен")
                return None
            else:
                print(f"   [OK] Запрос подходит для создания документа")
            
            if found_keywords or has_document_action or is_generating_document_text:
                print(f"\n{'='*60}")
                print(f"[START] НАЧИНАЕМ ПРОЦЕСС СОЗДАНИЯ ДОКУМЕНТА")
                print(f"   found_keywords: {found_keywords}")
                print(f"   has_document_action: {has_document_action}")
                print(f"   is_generating_document_text: {is_generating_document_text}")
                print(f"{'='*60}\n")
                
                # ПРИОРИТЕТ 1: Если обнаружили шаблон по названию + данные
                if detected_template_by_name:
                    print(f"[PRIORITY] ПРИОРИТЕТ 1: Используем обнаруженный шаблон: {detected_template_by_name['name']}")
                    template_recommendation = {
                        "suggested_template_id": detected_template_by_name['template_id'],
                        "suggested_template_name": detected_template_by_name['name'],
                        "confidence": 95,
                        "reasoning": f"Название шаблона обнаружено в запросе: '{detected_template_by_name['name']}'",
                        "document_category": self._get_category_from_name(detected_template_by_name['name'])
                    }
                    print(f"   Template ID: {template_recommendation['suggested_template_id']}")
                    print(f"   Category: {template_recommendation['document_category']}")
                    print(f"   Confidence: {template_recommendation['confidence']}%")
                else:
                    # ПРИОРИТЕТ 2: Пробуем найти шаблон по ключевым словам
                    print(f"\nПоиск шаблона по ключевым словам в запросе...")
                    detected_template = self._detect_template_from_request(user_messages_only, templates)
                    
                    if detected_template:
                        print(f" Найден шаблон напрямую: {detected_template['name']}")
                        template_recommendation = {
                            "suggested_template_id": detected_template['template_id'],
                            "suggested_template_name": detected_template['name'],
                            "confidence": 95,
                            "reasoning": f"Пользователь явно запросил шаблон '{detected_template['name']}'",
                            "document_category": self._get_category_from_name(detected_template['name'])
                        }
                    else:
                        # ПРИОРИТЕТ 3: Используем AI классификацию
                        print(f" Используем AI для классификации документа...")
                        template_recommendation = openai_service.classify_document_type(
                            conversation_text, 
                            templates
                        )
                
                # Анализируем шаблоны и определяем нужные поля
                template_analysis = self._analyze_templates_for_fields(templates)
                
                # НОВАЯ ФУНКЦИЯ: Расширенное извлечение данных
                required_fields = self._get_required_fields_for_templates(template_analysis)
                template_fields = list(required_fields.keys())
                
                advanced_extraction = openai_service.extract_structured_data_advanced(
                    conversation_text,
                    template_fields
                )
                
                # Преобразуем расширенные данные в стандартный формат
                user_info = {}
                for field_name, field_data in advanced_extraction.get("extracted_data", {}).items():
                    if isinstance(field_data, dict):
                        user_info[field_name] = field_data.get("value", "")
                    else:
                        user_info[field_name] = field_data
                
                # ПРОСТОЕ ИЗВЛЕЧЕНИЕ ДАННЫХ
                import re
                
                print(f"\nИЗВЛЕЧЕНИЕ ДАННЫХ ИЗ СООБЩЕНИЯ:")
                print(f"Анализируемый текст: '{conversation_text[:200]}...'")

                # Простые паттерны для извлечения данных (поддержка = и :)
                field_patterns = {
                    'last_name': r'(?:Фамилия|фамилия)[\s:=]+([А-ЯЁ][а-яё]+)',
                    'first_name': r'(?:Имя|имя)[\s:=]+([А-ЯЁа-яё]+)',
                    'middle_name': r'(?:Отчество|отчество)[\s:=]+([А-ЯЁ][а-яё]+)',
                    'inn': r'(?:ИНН|инн)[\s:=]+(\d{10,12})',
                    'phone': r'(?:Телефон|телефон)[\s:=]+(\+?[0-9\s\(\)\-]+)',
                    'email': r'(?:E-mail|email|почта)[\s:=]+([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                    'region': r'(?:Регион|регион)[\s:=]+([А-ЯЁа-яё\s\-]+?)(?:\s+Город|$)',
                    'city': r'(?:Город|город)[\s:=]+([А-ЯЁа-яё\-]+)',
                    'street': r'(?:Улица|улица)[\s:=]+([А-ЯЁа-яё\s\-]+?)(?:\s+Дом|$)',
                    'house': r'(?:Дом|дом)[\s:=]+(\d+[а-яА-Яa-zA-Z]*)',
                    'apartment': r'(?:Кв\.?|квартира)[\s:=]+(\d+)',
                    'birth_date': r'(?:Число[,\s]*месяц[,\s]*год\s+рождения|дата\s+рождения)[\s:=]+([^А-ЯЁ]+?)(?:\s+[А-ЯЁ]|$)',
                    'passport': r'(?:Паспортные\s+данные|паспорт)[\s:=]+([^А-ЯЁ][^\n]+?)(?:\s+[А-ЯЁ][а-яё]+:|$)',
                    'education': r'(?:Образование|образование)[\s:=]+([^\n]+?)(?:\s+Место\s+работы|$)',
                    'work_info': r'(?:Место\s+работы[,\s]*должность|работа)[\s:=]+([^\n]+?)(?:\s+Сфера|$)',
                    'activity_sphere': r'(?:Сфера\s+бизнеса|сфера)[\s:=]+([^\n]+?)(?:\s+Опыт\s+предпринимательской|$)',
                    'business_experience': r'(?:Опыт\s+предпринимательской\s+деятельности)[\s:=]+([^\n]+?)(?:\s+Опыт\s+общественной|$)',
                    'public_activity_experience': r'(?:Опыт\s+общественной\s+деятельности)[\s:=]+([^\n]+?)(?:\s+Эксперт|$)',
                    'expertise_area': r'(?:Эксперт\s+в\s+отрасли)[\s:=]+([^\n]+?)(?:\s+Выборные|$)',
                    'elected_position': r'(?:Выборные\s+должности)[\s:=]+([^\n]+?)(?:\s+Дополнения|$)',
                    'additional_info': r'(?:Дополнения)[\s:=]+([^\n]+?)(?:\s+Дата\s+заполнения|$)',
                }
                
                for field_name, pattern in field_patterns.items():
                    match = re.search(pattern, conversation_text, re.IGNORECASE)
                    if match:
                        user_info[field_name] = match.group(1).strip()
                        print(f"   + {field_name}: {match.group(1)[:50]}...")
                
                # Объединяем ФИО если есть отдельные поля
                if user_info.get('last_name') or user_info.get('first_name') or user_info.get('middle_name'):
                    full_name_parts = []
                    if user_info.get('last_name'):
                        full_name_parts.append(user_info['last_name'])
                    if user_info.get('first_name'):
                        full_name_parts.append(user_info['first_name'])
                    if user_info.get('middle_name'):
                        full_name_parts.append(user_info['middle_name'])
                    user_info['full_name'] = ' '.join(full_name_parts)
                    print(f"   + full_name (объединено): {user_info['full_name']}")
                
                # Объединяем адрес если есть отдельные поля
                if user_info.get('region') or user_info.get('city') or user_info.get('street') or user_info.get('house'):
                    address_parts = []
                    if user_info.get('region'):
                        address_parts.append(user_info['region'])
                    if user_info.get('city'):
                        address_parts.append(f"г. {user_info['city']}")
                    if user_info.get('street'):
                        address_parts.append(f"ул. {user_info['street']}")
                    if user_info.get('house'):
                        house_str = f"д. {user_info['house']}"
                        if user_info.get('apartment'):
                            house_str += f", кв. {user_info['apartment']}"
                        address_parts.append(house_str)
                    user_info['address'] = ', '.join(address_parts)
                    print(f"   + address (объединен): {user_info['address'][:80]}...")
                
                # Финальная проверка
                extracted_fields = [k for k, v in user_info.items() if v and v.strip()]
                print(f"\nИЗВЛЕЧЕНО ПОЛЕЙ: {len(extracted_fields)}")
                for field in extracted_fields:
                    print(f"   {field}: {user_info[field][:50]}...")
                
                # Базовые поля
                user_info.setdefault("user_id", user_id)
                user_info.setdefault("full_name", "")
                user_info.setdefault("email", "")
                user_info.setdefault("phone", "")
                user_info.setdefault("organization", "")
                user_info.setdefault("position", "")
                user_info.setdefault("inn", "")
                
                # НОВАЯ ФУНКЦИЯ: Анализ полноты данных
                completeness_analysis = openai_service.analyze_document_data_completeness(
                    user_info,
                    template_fields,
                    conversation_text
                )
                
                missing_fields = self._get_missing_fields(user_info, required_fields)
                
                # Формируем расширенный ответ
                result = {
                    "suggested": True,
                    "templates": templates,
                    "user_data": user_info,
                    "conversation_data": {
                        "message": message,
                        "response": response
                    },
                    "required_fields": required_fields,
                    "template_analysis": template_analysis,
                    # НОВЫЕ ПОЛЯ
                    "template_recommendation": template_recommendation,
                    "advanced_extraction": advanced_extraction,
                    "completeness_analysis": completeness_analysis,
                    "missing_fields": missing_fields,
                }
                
                # Определяем сообщение на основе анализа
                completeness_score = completeness_analysis.get("completeness_score", 0)
                confidence_score = completeness_analysis.get("confidence_score", 0)
                can_generate = completeness_analysis.get("can_generate", False)
                
                # СНИЖЕН ПОРОГ: создаем документ даже при 50% полноты (было 70%)
                # Это позволит создавать документы с частичными данными
                print(f"\n{'='*60}")
                print(f"[ANALYSIS] АНАЛИЗ ПОЛНОТЫ ДАННЫХ:")
                print(f"   completeness_score: {completeness_score}%")
                print(f"   confidence_score: {confidence_score}%")
                print(f"   can_generate: {can_generate}")
                print(f"   Порог для создания: 50%")
                
                # НОВАЯ ЛОГИКА: Если шаблон обнаружен явно по названию + есть структурированные данные
                # то создаем документ ВСЕГДА, независимо от оценки AI
                # ДОПОЛНИТЕЛЬНО: Для заявок на вступление создаем документ даже без структурированных данных
                is_application_request = ("заявк" in message_lower or "заявлен" in message_lower) and "вступлен" in message_lower
                force_create = (detected_template_by_name is not None and has_structured_data) or (detected_template_by_name is not None and is_application_request)
                print(f"   force_create (шаблон+данные): {force_create}")
                print(f"   Решение: {'[OK] СОЗДАЕМ' if (force_create or (can_generate and completeness_score >= 50)) else '[ERROR] НЕ создаем'}")
                print(f"{'='*60}\n")
                
                if force_create or (can_generate and completeness_score >= 50):
                    # Достаточно данных для создания документа - СОЗДАЁМ АВТОМАТИЧЕСКИ
                    suggested_template_id = template_recommendation.get("suggested_template_id")
                    suggested_template = template_recommendation.get("suggested_template_name", "")
                    
                    # FALLBACK: Если AI не определил шаблон, выбираем вручную по ключевым словам
                    if not suggested_template_id and templates:
                        print(" AI не определил шаблон, выбираем по ключевым словам")
                        doc_category = template_recommendation.get("document_category", "документ")
                        
                        # Ищем подходящий шаблон по категории
                        for template in templates:
                            template_name_lower = template.get('name', '').lower()
                            
                            if doc_category == "заявление" and "заявление" in template_name_lower:
                                suggested_template_id = template['template_id']
                                suggested_template = template['name']
                                print(f" Выбран шаблон: {suggested_template}")
                                break
                            elif doc_category == "анкета" and "анкета" in template_name_lower:
                                suggested_template_id = template['template_id']
                                suggested_template = template['name']
                                print(f" Выбран шаблон: {suggested_template}")
                                break
                            elif doc_category == "договор" and "договор" in template_name_lower:
                                suggested_template_id = template['template_id']
                                suggested_template = template['name']
                                print(f" Выбран шаблон: {suggested_template}")
                                break
                        
                        # Если все еще не нашли, берем первый шаблон
                        if not suggested_template_id and templates:
                            suggested_template_id = templates[0]['template_id']
                            suggested_template = templates[0]['name']
                            print(f" Используем первый доступный шаблон: {suggested_template}")
                    
                    # Если есть рекомендуемый шаблон, создаём документ автоматически
                    created_document = None
                    if suggested_template_id:
                        try:
                            print(f" Автоматическое создание документа: {suggested_template}")
                            created_document = self.create_document_from_template(
                                user_id=user_id,
                                template_id=suggested_template_id,
                                user_data=user_info,
                                conversation_data={
                                    "message": message,
                                    "response": response
                                },
                                send_email=True  # Включаем автоматическую отправку email
                            )
                            print(f" Документ создан: {created_document.get('filepath')}")
                            
                            # НОВОЕ: Сохраняем информацию о документе в Google Sheets
                            if created_document and created_document.get("status") == "success":
                                try:
                                    google_sheets_service.save_document({
                                        "user_id": user_id,
                                        "full_name": user_info.get("full_name", ""),
                                        "email": user_info.get("email", ""),
                                        "document_type": template_recommendation.get("document_category", "документ"),
                                        "template_name": suggested_template,
                                        "filepath": created_document.get("filepath", ""),
                                        "download_url": created_document.get("download_url", ""),
                                        "completeness_score": completeness_score,
                                        "confidence_score": confidence_score,
                                        "data_quality": completeness_analysis.get("data_quality", "")
                                    })
                                except Exception as e:
                                    print(f" Ошибка сохранения в Google Sheets (не критично): {e}")
                            
                        except Exception as e:
                            print(f" Ошибка автоматического создания: {e}")
                            import traceback
                            traceback.print_exc()
                    
                    # Формируем сообщение с результатом
                    if created_document and created_document.get("status") == "success":
                        # Документ успешно создан
                        download_url = created_document.get("download_url", "")
                        filename = created_document.get("filepath", "").split("/")[-1]
                        
                        result["message"] = f""" **Документ успешно создан!**

**Тип документа:** {template_recommendation.get('document_category', 'документ')}
**Шаблон:** {suggested_template}
**Имя файла:** {filename}

**Заполнено полей:** {len(completeness_analysis.get('filled_fields', []))} из {len(template_fields)}
**Оценка полноты:** {completeness_score}%
**Оценка уверенности:** {confidence_score}%
**Качество данных:** {completeness_analysis.get('data_quality', 'неизвестно')}

 **Ваш документ готов к скачиванию!**

**Скачать документ:** http://localhost:8000{download_url}"""
                        
                        result["created_document"] = created_document
                        result["needs_data"] = False
                    else:
                        # Не удалось создать автоматически, предлагаем выбрать вручную
                        result["message"] = f""" **Готов создать документ!**

**Тип документа:** {template_recommendation.get('document_category', 'документ')}
**Рекомендуемый шаблон:** {suggested_template or 'выберите из списка'}
**Уверенность:** {template_recommendation.get('confidence', 0)}%

**Заполнено полей:** {len(completeness_analysis.get('filled_fields', []))} из {len(template_fields)}
**Оценка полноты:** {completeness_score}%
**Оценка уверенности:** {confidence_score}%
**Качество данных:** {completeness_analysis.get('data_quality', 'неизвестно')}

 **У меня есть достаточно данных для создания документа.**
Выберите шаблон для создания готового документа:"""
                        result["needs_data"] = False
                else:
                    # Недостаточно данных по оценке AI
                    print(f"\n[WARNING] AI считает данных недостаточно, но проверяем force_create...")
                    print(f"   detected_template_by_name: {detected_template_by_name is not None}")
                    print(f"   has_structured_data: {has_structured_data}")
                    
                    # Если есть явный шаблон + структурированные данные - создаем в любом случае!
                    # ДОПОЛНИТЕЛЬНО: Для заявок на вступление создаем документ даже без структурированных данных
                    is_application_request_fallback = ("заявк" in message_lower or "заявлен" in message_lower) and "вступлен" in message_lower
                    if detected_template_by_name and (has_structured_data or is_application_request_fallback):
                        print(f"   [OK] FORCE CREATE: Создаем документ принудительно!")
                        suggested_template_id = detected_template_by_name['template_id']
                    else:
                        print(f"   [WARNING] Пытаемся найти шаблон другим способом...")
                        suggested_template_id = template_recommendation.get("suggested_template_id")
                    
                    # Fallback выбор шаблона
                    if not suggested_template_id and templates:
                        doc_category = template_recommendation.get("document_category", "документ")
                        for template in templates:
                            template_name_lower = template.get('name', '').lower()
                            if doc_category.lower() in template_name_lower or "заявление" in template_name_lower:
                                suggested_template_id = template['template_id']
                                break
                        if not suggested_template_id:
                            suggested_template_id = templates[0]['template_id']
                    
                    # СОЗДАЕМ ДОКУМЕНТ несмотря на неполные данные
                    created_document = None
                    if suggested_template_id:
                        try:
                            print(f" Создание документа с неполными данными ({completeness_score}%)")
                            created_document = self.create_document_from_template(
                                user_id=user_id,
                                template_id=suggested_template_id,
                                user_data=user_info,
                                conversation_data={
                                    "message": message,
                                    "response": response
                                },
                                send_email=True  # Включаем автоматическую отправку email
                            )
                            print(f" Документ создан с неполными данными: {created_document.get('filepath')}")
                            
                            # НОВОЕ: Сохраняем информацию о документе в Google Sheets
                            if created_document and created_document.get("status") == "success":
                                try:
                                    # Находим название шаблона
                                    template_name = next((t['name'] for t in templates if t['template_id'] == suggested_template_id), "Неизвестный шаблон")
                                    
                                    google_sheets_service.save_document({
                                        "user_id": user_id,
                                        "full_name": user_info.get("full_name", ""),
                                        "email": user_info.get("email", ""),
                                        "document_type": template_recommendation.get("document_category", "документ"),
                                        "template_name": template_name,
                                        "filepath": created_document.get("filepath", ""),
                                        "download_url": created_document.get("download_url", ""),
                                        "completeness_score": completeness_score,
                                        "confidence_score": confidence_score,
                                        "data_quality": completeness_analysis.get("data_quality", "")
                                    })
                                except Exception as e:
                                    print(f" Ошибка сохранения в Google Sheets (не критично): {e}")
                            
                        except Exception as e:
                            print(f" Ошибка создания: {e}")
                            import traceback
                            traceback.print_exc()
                    
                    # Формируем сообщение
                    missing_list = "\n".join([f"• {field}" for field in missing_fields])
                    
                    if created_document and created_document.get("status") == "success":
                        # Документ создан, но с предупреждением о неполных данных
                        download_url = created_document.get("download_url", "")
                        filename = created_document.get("filepath", "").split("/")[-1]
                        
                        result["message"] = f""" **Документ создан!**

 **Внимание:** Некоторые поля остались незаполненными ({completeness_score}% полноты).

**Отсутствуют данные:**
{missing_list}

**Рекомендация:** Отредактируйте документ вручную или предоставьте недостающие данные для создания нового.

 **Документ доступен для скачивания:**"""
                        
                        result["created_document"] = created_document
                        result["needs_data"] = False
                    else:
                        # Не удалось создать даже с неполными данными
                        recommendations = completeness_analysis.get("recommendations", [])
                        recommendations_text = "\n".join([f"- {rec}" for rec in recommendations])
                        
                        suggested_questions = completeness_analysis.get("suggested_questions", [])
                        questions_text = "\n".join([f"? {q}" for q in suggested_questions[:3]])
                        
                        result["message"] = f""" **Готов создать документ!**

**Тип документа:** {template_recommendation.get('document_category', 'документ')}

**Заполнено полей:** {len(completeness_analysis.get('filled_fields', []))} из {len(template_fields)}
**Оценка полноты:** {completeness_score}%

 **Отсутствуют данные:**
{missing_list}

{recommendations_text}

{questions_text if questions_text else ''}

**Предоставьте недостающие данные для создания документа.**"""
                        result["needs_data"] = True
                
                print(f"\n{'='*60}")
                print(f"[RESULT] ВОЗВРАЩАЕМ RESULT с document_suggestion:")
                print(f"   - suggested: {result.get('suggested')}")
                print(f"   - created_document exists: {result.get('created_document') is not None}")
                if result.get('created_document'):
                    print(f"   - created_document.status: {result.get('created_document', {}).get('status')}")
                    print(f"   - created_document.filepath: {result.get('created_document', {}).get('filepath', 'N/A')}")
                print(f"   - needs_data: {result.get('needs_data')}")
                print(f"   - message length: {len(result.get('message', ''))}")
                print(f"{'='*60}\n")
                return result
            
            print(f"\n[ERROR] Ключевые слова не найдены, возвращаем None")
            return None
            
        except Exception as e:
            print(f" ОШИБКА проверки создания документа: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _analyze_templates_for_fields(self, templates: List[Dict]) -> Dict:
        """
        Анализирует шаблоны и определяет, какие поля нужны для заполнения
        """
        try:
            field_analysis = {
                "common_fields": set(),
                "template_specific": {},
                "all_possible_fields": set()
            }
            
            # Возможные поля для поиска в шаблонах
            field_patterns = {
                "full_name": ["{{full_name}}", "{{name}}", "{{фио}}", "{{fio}}", "ФИО", "Имя", "Фамилия"],
                "email": ["{{email}}", "{{почта}}", "{{mail}}", "Email", "Почта", "Электронная почта"],
                "phone": ["{{phone}}", "{{телефон}}", "{{tel}}", "Телефон", "Контактный телефон"],
                "organization": ["{{organization}}", "{{организация}}", "{{org}}", "Организация", "Компания", "Предприятие"],
                "position": ["{{position}}", "{{должность}}", "{{pos}}", "Должность", "Позиция", "Роль"],
                "inn": ["{{inn}}", "{{инн}}", "ИНН", "Индивидуальный номер налогоплательщика"],
                "address": ["{{address}}", "{{адрес}}", "{{addr}}", "Адрес", "Место жительства"],
                "passport": ["{{passport}}", "{{паспорт}}", "{{pass}}", "Паспорт", "Документ"],
                "birth_date": ["{{birth_date}}", "{{дата_рождения}}", "{{birth}}", "Дата рождения", "Родился"],
                "business_type": ["{{business_type}}", "{{тип_бизнеса}}", "{{biz_type}}", "Тип бизнеса", "Вид деятельности"]
            }
            
            for template in templates:
                template_id = template.get('template_id', '')
                template_name = template.get('name', '')
                template_fields = set()
                
                # Анализируем название шаблона
                template_name_lower = template_name.lower()
                
                # Определяем тип документа по названию
                doc_type = self._determine_document_type(template_name_lower)
                
                # Базовые поля для каждого типа документа
                base_fields = self._get_base_fields_for_document_type(doc_type)
                template_fields.update(base_fields)
                
                # Добавляем поля в общий анализ
                field_analysis["common_fields"].update(template_fields)
                field_analysis["all_possible_fields"].update(template_fields)
                field_analysis["template_specific"][template_id] = {
                    "name": template_name,
                    "type": doc_type,
                    "fields": list(template_fields)
                }
            
            # Преобразуем множества в списки для JSON сериализации
            field_analysis["common_fields"] = list(field_analysis["common_fields"])
            field_analysis["all_possible_fields"] = list(field_analysis["all_possible_fields"])
            
            print(f"Анализ шаблонов: {field_analysis}")
            return field_analysis
            
        except Exception as e:
            print(f"Ошибка анализа шаблонов: {e}")
            return {"common_fields": [], "template_specific": {}, "all_possible_fields": []}
    
    def _determine_document_type(self, template_name: str) -> str:
        """
        Определяет тип документа по названию шаблона
        """
        if any(word in template_name for word in ["заявление", "заявка", "вступление"]):
            return "заявление"
        elif any(word in template_name for word in ["анкета", "анкетирование", "опрос"]):
            return "анкета"
        elif any(word in template_name for word in ["договор", "соглашение", "контракт"]):
            return "договор"
        elif any(word in template_name for word in ["жалоба", "претензия"]):
            return "жалоба"
        elif any(word in template_name for word in ["отчет", "отчетность"]):
            return "отчет"
        elif any(word in template_name for word in ["справка", "свидетельство"]):
            return "справка"
        elif any(word in template_name for word in ["протокол", "заседание"]):
            return "протокол"
        else:
            return "документ"
    
    def _get_base_fields_for_document_type(self, doc_type: str) -> set:
        """
        Возвращает базовые поля для конкретного типа документа
        """
        field_mapping = {
            "заявление": {"full_name", "email", "phone", "organization"},
            "анкета": {"full_name", "email", "phone", "organization", "position", "inn", "address", "birth_date"},
            "договор": {"full_name", "email", "phone", "organization", "inn", "address"},
            "жалоба": {"full_name", "email", "phone", "address"},
            "отчет": {"full_name", "organization", "position", "business_type"},
            "справка": {"full_name", "email", "phone", "address"},
            "протокол": {"full_name", "organization", "position"},
            "документ": {"full_name", "email", "phone", "organization"}
        }
        
        return field_mapping.get(doc_type, {"full_name", "email", "phone"})
    
    def _get_required_fields_for_templates(self, template_analysis: Dict) -> Dict[str, str]:
        """
        Определяет обязательные поля на основе анализа шаблонов
        """
        field_labels = {
            "full_name": "ФИО",
            "last_name": "Фамилия",
            "first_name": "Имя",
            "middle_name": "Отчество",
            "email": "Email",
            "phone": "Телефон",
            "organization": "Организация",
            "position": "Должность",
            "inn": "ИНН",
            "address": "Адрес",
            "passport": "Паспорт",
            "birth_date": "Дата рождения",
            "business_type": "Тип бизнеса",
            "region": "Регион",
            "city": "Город",
            "street": "Улица",
            "house": "Дом",
            "apartment": "Квартира",
            "education": "Образование",
            "work_info": "Информация о работе",
            "activity_sphere": "Сфера деятельности",
            "business_experience": "Опыт в бизнесе",
            "public_activity_experience": "Опыт общественной деятельности",
            "expertise_area": "Область экспертизы",
            "elected_position": "Выборная должность",
            "additional_info": "Дополнительная информация"
        }
        
        # Берем наиболее часто встречающиеся поля
        common_fields = template_analysis.get("common_fields", [])
        required_fields = {}
        
        for field in common_fields:
            if field in field_labels:
                required_fields[field] = field_labels[field]
        
        # Если нет общих полей, используем базовые
        if not required_fields:
            required_fields = {
                "full_name": "ФИО",
                "email": "Email",
                "phone": "Телефон",
                "organization": "Организация"
            }
        
        return required_fields
    
    def _analyze_template_content(self, template_id: str) -> Dict[str, str]:
        """
        Анализирует содержимое конкретного шаблона и извлекает поля из плейсхолдеров
        """
        try:
            from services import document_service
            import re
            import os
            
            # Получаем информацию о шаблоне
            templates = document_service.get_templates_list()
            template_info = None
            
            for template in templates:
                if template['template_id'] == template_id:
                    template_info = template
                    break
            
            if not template_info:
                print(f"Шаблон {template_id} не найден")
                return self._get_default_fields_for_template(template_id)
            
            template_name = template_info['name']
            template_path = template_info.get('file_path', '')
            
            print(f"Анализируем шаблон: {template_name} (путь: {template_path})")
            
            # Извлекаем плейсхолдеры из файла
            placeholders = self._extract_placeholders_from_file(template_path)
            
            if placeholders:
                print(f"Найдены плейсхолдеры: {list(placeholders.keys())}")
                return placeholders
            else:
                print("Плейсхолдеры не найдены, используем анализ по названию")
                # Fallback к анализу по названию
                template_name_lower = template_name.lower()
            
            # Определяем поля на основе названия и типа шаблона
            if 'жалоба' in template_name_lower or 'complaint' in template_name_lower:
                return {
                    "full_name": "ФИО заявителя",
                    "email": "Email для связи",
                    "phone": "Телефон",
                    "organization": "Организация",
                    "inn": "ИНН организации",
                    "address": "Адрес организации",
                    "business_type": "Вид деятельности"
                }
            elif 'протокол' in template_name_lower or 'protocol' in template_name_lower:
                return {
                    "full_name": "ФИО участника",
                    "organization": "Организация",
                    "position": "Должность",
                    "email": "Email",
                    "phone": "Телефон"
                }
            elif 'договор' in template_name_lower or 'contract' in template_name_lower:
                return {
                    "full_name": "ФИО",
                    "organization": "Организация",
                    "inn": "ИНН",
                    "address": "Юридический адрес",
                    "email": "Email",
                    "phone": "Телефон",
                    "position": "Должность"
                }
            elif 'заявка' in template_name_lower or 'application' in template_name_lower:
                return {
                    "full_name": "ФИО заявителя",
                    "organization": "Название организации",
                    "inn": "ИНН",
                    "email": "Email",
                    "phone": "Телефон",
                    "business_type": "Вид деятельности",
                    "address": "Адрес"
                }
            elif 'россия' in template_name_lower or 'russia' in template_name_lower:
                # Специальные поля для документов Опоры России
                return {
                    "full_name": "ФИО",
                    "email": "Email",
                    "phone": "Телефон",
                    "organization": "Название организации",
                    "inn": "ИНН организации",
                    "address": "Адрес организации",
                    "business_type": "Вид деятельности",
                    "position": "Должность"
                }
            elif 'вступление' in template_name_lower or 'membership' in template_name_lower:
                # Поля для заявки на вступление
                return {
                    "full_name": "ФИО заявителя",
                    "organization": "Название организации",
                    "inn": "ИНН",
                    "email": "Email",
                    "phone": "Телефон",
                    "business_type": "Вид деятельности",
                    "address": "Адрес организации",
                    "position": "Должность"
                }
            elif '73' in template_name_lower or 'федеральный' in template_name_lower:
                # Поля для документов по 73-ФЗ
                return {
                    "full_name": "ФИО",
                    "organization": "Организация",
                    "inn": "ИНН",
                    "email": "Email",
                    "phone": "Телефон",
                    "address": "Адрес",
                    "business_type": "Вид деятельности"
                }
            else:
                # Для неизвестных типов используем базовые поля
                return self._get_default_fields_for_template(template_id)
                
        except Exception as e:
            print(f"Ошибка анализа шаблона {template_id}: {e}")
            return self._get_default_fields_for_template(template_id)
    
    def _extract_placeholders_from_file(self, file_path: str) -> Dict[str, str]:
        """
        Извлекает плейсхолдеры из файла шаблона
        """
        try:
            import re
            import os
            from docx import Document
            
            if not os.path.exists(file_path):
                print(f"Файл не найден: {file_path}")
                return {}
            
            placeholders = {}
            content = ""
            
            # Определяем тип файла
            if file_path.lower().endswith('.docx'):
                # Читаем DOCX файл
                try:
                    doc = Document(file_path)
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                content += cell.text + " "
                except Exception as e:
                    print(f"Ошибка чтения DOCX файла: {e}")
                    return {}
            elif file_path.lower().endswith('.txt'):
                # Читаем текстовый файл
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except Exception as e:
                    print(f"Ошибка чтения TXT файла: {e}")
                    return {}
            else:
                print(f"Неподдерживаемый тип файла: {file_path}")
                return {}
            
            # Ищем плейсхолдеры в формате {{field_name}}
            placeholder_pattern = r'\{\{([^}]+)\}\}'
            matches = re.findall(placeholder_pattern, content)
            
            # Маппинг плейсхолдеров на человекочитаемые названия
            field_labels = {
                "full_name": "ФИО",
                "name": "ФИО",
                "фио": "ФИО",
                "fio": "ФИО",
                "last_name": "Фамилия",
                "фамилия": "Фамилия",
                "first_name": "Имя",
                "имя": "Имя",
                "middle_name": "Отчество",
                "отчество": "Отчество",
                "email": "Email",
                "почта": "Email",
                "mail": "Email",
                "phone": "Телефон",
                "телефон": "Телефон",
                "tel": "Телефон",
                "organization": "Организация",
                "организация": "Организация",
                "org": "Организация",
                "position": "Должность",
                "должность": "Должность",
                "pos": "Должность",
                "inn": "ИНН",
                "инн": "ИНН",
                "address": "Адрес",
                "адрес": "Адрес",
                "addr": "Адрес",
                "passport": "Паспорт",
                "паспорт": "Паспорт",
                "pass": "Паспорт",
                "birth_date": "Дата рождения",
                "дата_рождения": "Дата рождения",
                "birth": "Дата рождения",
                "business_type": "Тип бизнеса",
                "тип_бизнеса": "Тип бизнеса",
                "biz_type": "Тип бизнеса",
                "region": "Регион",
                "регион": "Регион",
                "city": "Город",
                "город": "Город",
                "street": "Улица",
                "улица": "Улица",
                "house": "Дом",
                "дом": "Дом",
                "apartment": "Квартира",
                "квартира": "Квартира",
                "education": "Образование",
                "образование": "Образование",
                "work_info": "Информация о работе",
                "информация_о_работе": "Информация о работе",
                "activity_sphere": "Сфера деятельности",
                "сфера_деятельности": "Сфера деятельности",
                "business_experience": "Опыт в бизнесе",
                "опыт_в_бизнесе": "Опыт в бизнесе",
                "public_activity_experience": "Опыт общественной деятельности",
                "опыт_общественной_деятельности": "Опыт общественной деятельности",
                "expertise_area": "Область экспертизы",
                "область_экспертизы": "Область экспертизы",
                "elected_position": "Выборная должность",
                "выборная_должность": "Выборная должность",
                "additional_info": "Дополнительная информация",
                "дополнительная_информация": "Дополнительная информация",
                "user_id": "ID пользователя",
                "date": "Дата",
                "дата": "Дата",
                "time": "Время",
                "время": "Время"
            }
            
            # Обрабатываем найденные плейсхолдеры
            for match in matches:
                field_name = match.strip().lower()
                if field_name in field_labels:
                    placeholders[field_name] = field_labels[field_name]
                else:
                    # Если поле не найдено в маппинге, используем оригинальное название
                    placeholders[field_name] = match.strip()
            
            print(f"Извлечено плейсхолдеров: {len(placeholders)}")
            for field, label in placeholders.items():
                print(f"  • {field} → {label}")
            
            return placeholders
            
        except Exception as e:
            print(f"Ошибка извлечения плейсхолдеров: {e}")
            return {}
    
    def _get_default_fields_for_template(self, template_id: str) -> Dict[str, str]:
        """
        Возвращает базовые поля для шаблона
        """
        return {
            "full_name": "ФИО",
            "email": "Email",
            "phone": "Телефон",
            "organization": "Организация"
        }
    
    def _analyze_data_completeness(self, user_data: Dict, required_fields: Dict[str, str]) -> Dict:
        """
        Анализирует полноту данных пользователя
        
        Args:
            user_data: Данные пользователя
            required_fields: Обязательные поля
        
        Returns:
            Анализ полноты данных
        """
        filled_fields = []
        missing_fields = []
        
        for field, label in required_fields.items():
            if user_data.get(field) and str(user_data[field]).strip():
                filled_fields.append(label)
            else:
                missing_fields.append(label)
        
        total_fields = len(required_fields)
        filled_count = len(filled_fields)
        completeness_score = int((filled_count / total_fields) * 100) if total_fields > 0 else 0
        
        # Определяем качество данных
        if completeness_score >= 90:
            data_quality = "excellent"
        elif completeness_score >= 70:
            data_quality = "good"
        elif completeness_score >= 50:
            data_quality = "fair"
        else:
            data_quality = "poor"
        
        # Генерируем рекомендуемые вопросы
        suggested_questions = []
        for field, label in required_fields.items():
            if not user_data.get(field) or not str(user_data[field]).strip():
                if field == "full_name":
                    suggested_questions.append("Как вас зовут? (ФИО в родительном падеже)")
                elif field == "email":
                    suggested_questions.append("Укажите ваш email адрес")
                elif field == "phone":
                    suggested_questions.append("Укажите ваш номер телефона")
                elif field == "organization":
                    suggested_questions.append("В какой организации вы работаете?")
                elif field == "position":
                    suggested_questions.append("Какую должность вы занимаете?")
                elif field == "inn":
                    suggested_questions.append("Укажите ИНН организации")
                elif field == "address":
                    suggested_questions.append("Укажите ваш адрес")
                elif field == "passport":
                    suggested_questions.append("Укажите серию и номер паспорта")
                elif field == "birth_date":
                    suggested_questions.append("Укажите дату рождения")
                elif field == "business_type":
                    suggested_questions.append("Какой тип бизнеса вы ведете?")
        
        # Генерируем рекомендации
        recommendations = []
        if completeness_score < 50:
            recommendations.append("Рекомендуется заполнить основные поля: ФИО, email, телефон")
        if "email" in missing_fields:
            recommendations.append("Email необходим для отправки готового документа")
        if "full_name" in missing_fields:
            recommendations.append("ФИО является обязательным полем для большинства документов")
        
        return {
            "completeness_score": completeness_score,
            "confidence_score": min(completeness_score + 10, 100),  # Уверенность немного выше полноты
            "data_quality": data_quality,
            "filled_fields": filled_fields,
            "missing_fields": missing_fields,
            "suggested_questions": suggested_questions[:5],  # Максимум 5 вопросов
            "recommendations": recommendations
        }
    
    def _get_missing_fields(self, user_data: Dict, required_fields: Dict[str, str]) -> List[str]:
        """
        Определяет, какие поля отсутствуют у пользователя
        """
        missing = []
        for field, label in required_fields.items():
            if not user_data.get(field):
                missing.append(label)
        return missing
    
    def _format_required_fields(self, required_fields: Dict[str, str]) -> str:
        """
        Форматирует список обязательных полей для отображения
        """
        if not required_fields:
            return "• Базовые контактные данные"
        
        formatted = []
        for field, label in required_fields.items():
            formatted.append(f"• **{label}**")
        
        return "\n".join(formatted)
    
    def preview_document(self, template_id: str, user_data: Dict) -> Dict:
        """
        Генерирует предпросмотр документа перед созданием
        
        Args:
            template_id: ID шаблона
            user_data: Данные пользователя
        
        Returns:
            Предпросмотр документа
        """
        try:
            # Получаем информацию о шаблоне
            templates = document_service.get_templates_list()
            template = next((t for t in templates if t['template_id'] == template_id), None)
            
            if not template:
                return {
                    "status": "error",
                    "message": "Шаблон не найден"
                }
            
            # Генерируем предпросмотр с помощью AI
            preview_text = openai_service.generate_document_preview(
                template.get('name', 'Документ'),
                user_data
            )
            
            # Анализируем полноту данных для этого шаблона
            template_fields = self._get_base_fields_for_document_type(
                self._determine_document_type(template.get('name', '').lower())
            )
            
            completeness = openai_service.analyze_document_data_completeness(
                user_data,
                list(template_fields),
                ""
            )
            
            return {
                "status": "success",
                "template_name": template.get('name', ''),
                "preview": preview_text,
                "completeness": completeness,
                "can_generate": completeness.get('can_generate', False)
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка генерации предпросмотра: {str(e)}"
            }
    
    def create_document_from_template(self, user_id: str, template_id: str, 
                                     user_data: Dict, conversation_data: Dict = None,
                                     send_email: bool = False) -> Dict:
        """
        Создать документ из шаблона (РАСШИРЕННАЯ ВЕРСИЯ)
        
        Args:
            user_id: ID пользователя
            template_id: ID шаблона
            user_data: Данные пользователя
            conversation_data: Данные из разговора
            send_email: Отправить документ на email
        
        Returns:
            Результат создания документа
        """
        try:
            # Дополняем данные пользователя
            user_data['user_id'] = user_id
            
            # Создаем документ
            filepath = document_service.fill_uploaded_template(
                template_id, 
                user_data, 
                conversation_data
            )
            
            # Получаем информацию о шаблоне
            templates = document_service.get_templates_list()
            template = next((t for t in templates if t['template_id'] == template_id), None)
            template_name = template.get('name', 'Документ') if template else 'Документ'
            
            result = {
                "status": "success",
                "filepath": filepath,
                "download_url": f"/api/documents/download?file={filepath}",
                "template_name": template_name,
                "message": f" Документ '{template_name}' успешно создан!"
            }
            
            # Отправляем на email если запрошено
            if send_email and user_data.get('email'):
                try:
                    from integrations import gmail_service
                    
                    # Инициализируем Gmail сервис если нужно
                    if not gmail_service.service:
                        gmail_service._initialize_service()
                    
                    # Отправляем email с ссылкой на скачивание вместо вложения
                    download_url = f"http://localhost:8000{result['download_url']}"
                    email_body = f"""
Здравствуйте!

Ваш документ '{template_name}' успешно создан!

Ссылка для скачивания: {download_url}

С уважением,
Команда "Опора России"
"""
                    
                    email_result = gmail_service.send_email(
                        to_email=user_data['email'],
                        subject=f"Созданный документ: {template_name}",
                        body=email_body
                    )
                    
                    if email_result:
                        result["message"] += f"\n[EMAIL] Ссылка на документ отправлена на {user_data['email']}"
                    else:
                        result["message"] += "\n[EMAIL] Документ создан, но не удалось отправить ссылку на email."
                except Exception as e:
                    result["message"] += f"\n Документ создан, но ошибка отправки email: {str(e)}"
            elif send_email and not user_data.get('email'):
                result["message"] += "\n[EMAIL] Документ создан, но email не указан для отправки."
            
            return result
            
        except Exception as e:
            return {
                "status": "error",
                "message": f" Ошибка создания документа: {str(e)}"
            }
    
    def _detect_template_from_request(self, user_message: str, templates: List[Dict]) -> Optional[Dict]:
        """
        Пытается найти шаблон напрямую по ключевым словам в запросе пользователя
        
        Args:
            user_message: Сообщение пользователя (lowercase)
            templates: Список доступных шаблонов
        
        Returns:
            Найденный шаблон или None
        """
        message_lower = user_message.lower()
        
        # Убираем частые слова для лучшего поиска
        stopwords = ["создать", "создайте", "заполнить", "заполните", "оформить", 
                     "документ", "шаблон", "нужен", "нужна", "для", "по", "в", "на"]
        
        # Извлекаем ключевые слова из запроса
        words = message_lower.split()
        keywords = [w for w in words if w not in stopwords and len(w) > 2]
        
        print(f"   Ключевые слова из запроса: {keywords}")
        
        # Проходим по всем шаблонам
        best_match = None
        best_score = 0
        
        for template in templates:
            template_name = template.get('name', '').lower()
            template_desc = template.get('description', '').lower()
            combined = f"{template_name} {template_desc}"
            
            # Подсчитываем совпадения слов
            score = 0
            matched_words = []
            
            for keyword in keywords:
                # Точное совпадение слова
                if keyword in combined:
                    score += 2
                    matched_words.append(keyword)
                # Частичное совпадение (начало слова)
                elif any(word.startswith(keyword) for word in combined.split()):
                    score += 1
                    matched_words.append(f"{keyword}*")
            
            # Бонус за совпадение с началом названия
            if keywords and template_name.startswith(keywords[0]):
                score += 3
            
            print(f"   - '{template.get('name')}': score={score}, matched={matched_words}")
            
            if score > best_score:
                best_score = score
                best_match = template
        
        # Требуем минимум 2 балла для считывания совпадения
        if best_score >= 2:
            return best_match
        
        return None
    
    def _get_category_from_name(self, template_name: str) -> str:
        """
        Определяет категорию документа по его названию
        
        Args:
            template_name: Название шаблона
        
        Returns:
            Категория документа
        """
        name_lower = template_name.lower()
        
        if "заявление" in name_lower or "заявка" in name_lower:
            return "заявление"
        elif "анкета" in name_lower:
            return "анкета"
        elif "договор" in name_lower:
            return "договор"
        elif "жалоба" in name_lower:
            return "жалоба"
        elif "протокол" in name_lower:
            return "протокол"
        elif "отчет" in name_lower or "отчёт" in name_lower:
            return "отчет"
        elif "справка" in name_lower:
            return "справка"
        else:
            return "другое"
    
    def get_conversation_history(self, user_id: str) -> List[Dict]:
        """
        Получить историю диалога пользователя
        
        Args:
            user_id: ID пользователя
        
        Returns:
            История диалога
        """
        return self.conversations.get(user_id, [])
    
    def clear_conversation(self, user_id: str):
        """
        Очистить историю диалога
        
        Args:
            user_id: ID пользователя
        """
        if user_id in self.conversations:
            del self.conversations[user_id]
    
    def _check_and_save_complaint(self, user_id: str, message: str, conversation_history: List[Dict]):
        """
        Проверить, является ли сообщение жалобой, и сохранить её
        
        Args:
            user_id: ID пользователя
            message: Сообщение пользователя
            conversation_history: История диалога
        """
        try:
            message_lower = message.lower()
            
            # Ключевые слова для распознавания жалоб
            complaint_keywords = [
                "жалоба", "жалуюсь", "недоволен", "недовольна", "проблема", "проблемы",
                "нарушение", "нарушения", "неправильно", "плохо", "ужасно", "кошмар",
                "безобразие", "возмущен", "возмущена", "возмущаюсь", "неприемлемо",
                "несправедливо", "обидно", "обижен", "обижена", "недовольство",
                "претензия", "претензии", "некачественно", "некачественный",
                "обман", "обманули", "обманывают", "мошенничество", "мошенники",
                "воровство", "воруют", "кража", "украли", "нечестно", "нечестные",
                "подделка", "подделки", "фальшивка", "фальшивые", "незаконно",
                "незаконные", "противозаконно", "нарушают закон", "нарушение закона",
                "дискриминация", "дискриминируют", "ущемляют права", "нарушают права",
                "не выполняют", "не выполняет", "не соблюдают", "не соблюдает",
                "игнорируют", "игнорирует", "не реагируют", "не реагирует",
                "не отвечают", "не отвечает", "не помогают", "не помогает",
                "отказываются", "отказывается", "отказ", "отказы", "отклонили",
                "отклонили заявку", "отклонили обращение", "не рассмотрели",
                "не рассмотрели заявку", "не рассмотрели обращение", "затягивают",
                "затягивают рассмотрение", "медленно работают", "медленно работает",
                "неэффективно", "неэффективная работа", "плохая работа", "плохо работают",
                "некомпетентно", "некомпетентные", "не знают", "не знает", "не понимают",
                "не понимает", "не разбираются", "не разбирается", "не умеют", "не умеет"
            ]
            
            # Проверяем наличие ключевых слов жалоб
            is_complaint = any(keyword in message_lower for keyword in complaint_keywords)
            
            if is_complaint:
                print(f"\n{'='*60}")
                print(f"[COMPLAINT DETECTED] Обнаружена жалоба!")
                print(f"User ID: {user_id}")
                print(f"Message: {message[:100]}...")
                print(f"{'='*60}\n")
                
                # Извлекаем данные пользователя из истории
                user_data = self._extract_user_data_from_history(conversation_history)
                
                # Определяем категорию жалобы
                category = self._determine_complaint_category(message_lower)
                
                # Определяем приоритет жалобы
                priority = self._determine_complaint_priority(message_lower)
                
                # Создаем жалобу
                complaint_data = {
                    "complaint_id": f"COMP_{user_id}_{int(datetime.now().timestamp())}",
                    "user_id": user_id,
                    "full_name": user_data.get("full_name", ""),
                    "email": user_data.get("email", ""),
                    "phone": user_data.get("phone", ""),
                    "organization": user_data.get("organization", ""),
                    "complaint_text": message,
                    "category": category,
                    "priority": priority
                }
                
                # Сохраняем жалобу в Google Sheets
                success = google_sheets_service.save_complaint(complaint_data)
                
                if success:
                    print(f"[SUCCESS] Жалоба сохранена в Google Sheets!")
                    print(f"   - ID: {complaint_data['complaint_id']}")
                    print(f"   - Категория: {category}")
                    print(f"   - Приоритет: {priority}")
                else:
                    print(f"[ERROR] Ошибка сохранения жалобы в Google Sheets")
                    
        except Exception as e:
            print(f"[ERROR] Ошибка при обработке жалобы: {e}")
            import traceback
            traceback.print_exc()
    
    def _extract_user_data_from_history(self, conversation_history: List[Dict]) -> Dict:
        """
        Извлечь данные пользователя из истории диалога
        
        Args:
            conversation_history: История диалога
        
        Returns:
            Словарь с данными пользователя
        """
        user_data = {}
        
        # Собираем все сообщения пользователя
        user_messages = []
        for msg in conversation_history:
            if msg.get("role") == "user":
                user_messages.append(msg.get("content", ""))
        
        # Объединяем все сообщения
        full_text = " ".join(user_messages)
        
        # Извлекаем данные регулярными выражениями
        import re
        
        patterns = {
            'full_name': r'(?:ФИО|фио|зовут|имя)[\s:=]+([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)',
            'email': r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            'phone': r'(\+?[0-9\s\(\)\-]{10,})',
            'organization': r'(?:организация|компания|ооо|общество)[\s:=]+([^,.\n]+)',
        }
        
        for field, pattern in patterns.items():
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                user_data[field] = match.group(1).strip()
        
        return user_data
    
    def _determine_complaint_category(self, message_lower: str) -> str:
        """
        Определить категорию жалобы
        
        Args:
            message_lower: Сообщение в нижнем регистре
        
        Returns:
            Категория жалобы
        """
        if any(word in message_lower for word in ["обман", "мошенничество", "воровство", "кража", "подделка"]):
            return "Мошенничество"
        elif any(word in message_lower for word in ["незаконно", "нарушение закона", "противозаконно"]):
            return "Нарушение закона"
        elif any(word in message_lower for word in ["дискриминация", "ущемляют права", "нарушают права"]):
            return "Нарушение прав"
        elif any(word in message_lower for word in ["некачественно", "плохая работа", "неэффективно"]):
            return "Качество услуг"
        elif any(word in message_lower for word in ["не отвечают", "игнорируют", "не реагируют"]):
            return "Отсутствие реакции"
        elif any(word in message_lower for word in ["отклонили", "отказ", "не рассмотрели"]):
            return "Отказ в рассмотрении"
        elif any(word in message_lower for word in ["затягивают", "медленно", "долго"]):
            return "Затягивание сроков"
        else:
            return "Общая"
    
    def _determine_complaint_priority(self, message_lower: str) -> str:
        """
        Определить приоритет жалобы
        
        Args:
            message_lower: Сообщение в нижнем регистре
        
        Returns:
            Приоритет жалобы
        """
        high_priority_keywords = [
            "срочно", "немедленно", "критично", "критическая", "опасно", "опасная",
            "мошенничество", "воровство", "кража", "обман", "незаконно", "противозаконно",
            "дискриминация", "ущемляют права", "нарушают права", "неприемлемо"
        ]
        
        if any(keyword in message_lower for keyword in high_priority_keywords):
            return "Высокий"
        elif any(word in message_lower for word in ["важно", "серьезно", "серьезная"]):
            return "Средний"
        else:
            return "Низкий"
    
    def start_interactive_autofill(self, user_id: str, message: str = None) -> Dict:
        """
        Начать интерактивный процесс автозаполнения документа
        
        Args:
            user_id: ID пользователя
            message: Сообщение пользователя (опционально)
        
        Returns:
            Начало интерактивного процесса
        """
        try:
            # Получаем список доступных шаблонов
            templates = document_service.get_templates_list()
            
            if not templates:
                return {
                    "status": "error",
                    "message": "Нет доступных шаблонов документов"
                }
            
            # Создаем сессию автозаполнения
            session_id = f"autofill_{user_id}_{int(datetime.now().timestamp())}"
            
            # Преобразуем шаблоны в правильный формат
            formatted_templates = []
            for template in templates:
                formatted_templates.append({
                    "template_id": template['template_id'],
                    "id": template['template_id'],
                    "name": template['name'],
                    "description": template.get('description', ''),
                    "category": self._get_category_from_name(template['name'])
                })
            
            self.autofill_sessions[session_id] = {
                "user_id": user_id,
                "status": "document_selection",
                "templates": formatted_templates,
                "selected_document": None,
                "user_data": {},
                "questions_asked": [],
                "questions_answered": [],
                "created_at": datetime.now().isoformat()
            }
            
            print(f"[DEBUG] Created session: {session_id} for user: {user_id}")
            print(f"[DEBUG] Total sessions: {len(self.autofill_sessions)}")
            
            # Формируем список документов для выбора (используем уже отформатированные шаблоны)
            document_list = formatted_templates
            
            return {
                "status": "success",
                "session_id": session_id,
                "message": "Выберите документ для автозаполнения:",
                "templates": document_list,
                "documents": document_list,  # Для обратной совместимости
                "next_step": "document_selection"
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка начала автозаполнения: {str(e)}"
            }
    
    def analyze_document_for_autofill(self, user_id: str, document_name: str) -> Dict:
        """
        Анализирует указанный документ и определяет нужные поля
        
        Args:
            user_id: ID пользователя
            document_name: Название документа для анализа
        
        Returns:
            Анализ документа и список нужных полей
        """
        try:
            # Находим сессию пользователя
            session = self._find_user_session(user_id)
            if not session:
                return {
                    "status": "error",
                    "message": "Сессия автозаполнения не найдена. Начните процесс заново."
                }
            
            # Находим выбранный документ
            print(f"[DEBUG] Ищем документ: {document_name}")
            print(f"[DEBUG] Доступные шаблоны в сессии: {len(session['templates'])}")
            for i, template in enumerate(session['templates']):
                print(f"[DEBUG]   {i+1}. {template.get('name', 'N/A')} (ID: {template.get('id', 'N/A')}, template_id: {template.get('template_id', 'N/A')})")
            
            selected_template = None
            for template in session['templates']:
                if template['name'].lower() == document_name.lower():
                    selected_template = template
                    print(f"[DEBUG] Найден по имени: {template['name']}")
                    break
            
            # Если не найден по имени, попробуем найти по ID
            if not selected_template:
                for template in session['templates']:
                    if template.get('template_id') == document_name or template.get('id') == document_name:
                        selected_template = template
                        print(f"[DEBUG] Найден по ID: {template.get('id')}")
                        break
            
            if not selected_template:
                return {
                    "status": "error",
                    "message": f"Документ '{document_name}' не найден"
                }
            
            # Обновляем сессию
            session['selected_document'] = selected_template
            session['status'] = 'analysis_complete'
            
            # Анализируем конкретный шаблон и определяем нужные поля
            required_fields = self._analyze_template_content(selected_template['template_id'])
            
            # Определяем категорию документа
            document_category = self._get_category_from_name(selected_template['name'])
            
            # Анализируем полноту данных
            completeness_analysis = self._analyze_data_completeness(session['user_data'], required_fields)
            
            return {
                "status": "success",
                "template_id": selected_template['template_id'],
                "document": {
                    "id": selected_template['template_id'],
                    "name": selected_template['name'],
                    "category": document_category
                },
                "required_fields": required_fields,
                "field_count": len(required_fields),
                "user_data": session['user_data'],
                "needs_data": completeness_analysis['completeness_score'] < 100,
                "completeness_analysis": completeness_analysis,
                "message": f"Документ '{selected_template['name']}' проанализирован. Требуется заполнить {len(required_fields)} полей.",
                "next_step": "data_collection"
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка анализа документа: {str(e)}"
            }
    
    def ask_questions_for_autofill(self, user_id: str, document_name: str, current_data: dict = None) -> Dict:
        """
        Задает вопросы для заполнения недостающих данных
        
        Args:
            user_id: ID пользователя
            document_name: Название документа
            current_data: Текущие данные пользователя
        
        Returns:
            Список вопросов для заполнения
        """
        try:
            # Находим сессию пользователя
            session = self._find_user_session(user_id)
            if not session:
                return {
                    "status": "error",
                    "message": "Сессия автозаполнения не найдена"
                }
            
            # Обновляем данные пользователя если предоставлены
            if current_data:
                session['user_data'].update(current_data)
            
            # Анализируем какие поля уже заполнены
            selected_template = session['selected_document']
            if not selected_template:
                return {
                    "status": "error",
                    "message": "Документ не выбран"
                }
            
            template_analysis = self._analyze_templates_for_fields([selected_template])
            required_fields = self._get_required_fields_for_templates(template_analysis)
            
            # Определяем недостающие поля
            missing_fields = []
            for field, label in required_fields.items():
                if not session['user_data'].get(field):
                    missing_fields.append((field, label))
            
            if not missing_fields:
                return {
                    "status": "success",
                    "message": "Все необходимые данные уже собраны!",
                    "questions": [],
                    "next_step": "ready_to_create"
                }
            
            # Генерируем вопросы для недостающих полей
            questions = []
            for i, (field, label) in enumerate(missing_fields[:5]):  # Максимум 5 вопросов за раз
                question_id = f"q_{field}_{i}"
                question_text = self._generate_question_for_field(field, label, selected_template['name'])
                
                questions.append({
                    "id": question_id,
                    "field": field,
                    "label": label,
                    "question": question_text,
                    "type": self._get_field_input_type(field)
                })
            
            # Обновляем статус сессии
            session['status'] = 'collecting_data'
            session['questions_asked'].extend([q['id'] for q in questions])
            
            return {
                "status": "success",
                "questions": questions,
                "progress": {
                    "total_fields": len(required_fields),
                    "filled_fields": len(required_fields) - len(missing_fields),
                    "remaining_fields": len(missing_fields)
                },
                "message": f"Пожалуйста, ответьте на {len(questions)} вопросов для заполнения документа.",
                "next_step": "answering_questions"
            }
            
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка генерации вопросов: {str(e)}"
            }
    
    def answer_question_for_autofill(self, user_id: str, question_id: str, answer: str) -> Dict:
        """
        Обрабатывает ответ пользователя на вопрос
        
        Args:
            user_id: ID пользователя
            question_id: ID вопроса
            answer: Ответ пользователя
        
        Returns:
            Результат обработки ответа
        """
        try:
            print(f"[DEBUG] answer_question_for_autofill called:")
            print(f"  user_id: {user_id}")
            print(f"  question_id: {question_id}")
            print(f"  answer: {answer}")
            print(f"  active sessions: {len(self.autofill_sessions)}")
            for sid, sess in self.autofill_sessions.items():
                print(f"    {sid}: user={sess['user_id']}, status={sess['status']}")
            
            # Находим сессию пользователя
            session = self._find_user_session(user_id)
            if not session:
                return {
                    "status": "error",
                    "message": f"Сессия автозаполнения не найдена для пользователя {user_id}. Активных сессий: {len(self.autofill_sessions)}"
                }
            
            # Извлекаем поле из question_id (формат: q_{field}_{i})
            if '_' in question_id and question_id.count('_') >= 2:
                field = question_id.split('_')[1]
            else:
                return {
                    "status": "error",
                    "message": f"Неверный формат question_id: {question_id}"
                }
            
            print(f"[DEBUG] Extracted field: {field}")
            print(f"[DEBUG] Session data: {session}")
            
            # Сохраняем ответ
            if 'user_data' not in session:
                session['user_data'] = {}
            if 'questions_answered' not in session:
                session['questions_answered'] = []
                
            session['user_data'][field] = answer
            session['questions_answered'].append(question_id)
            
            # Проверяем, все ли вопросы отвечены
            print(f"[DEBUG] questions_asked: {session.get('questions_asked', [])}")
            print(f"[DEBUG] questions_answered: {session.get('questions_answered', [])}")
            
            questions_asked = session.get('questions_asked', [])
            questions_answered = session.get('questions_answered', [])
            remaining_questions = set(questions_asked) - set(questions_answered)
            
            if not remaining_questions:
                # Все вопросы отвечены, проверяем полноту данных
                selected_template = session['selected_document']
                template_analysis = self._analyze_templates_for_fields([selected_template])
                required_fields = self._get_required_fields_for_templates(template_analysis)
                
                missing_fields = []
                for field, label in required_fields.items():
                    if not session['user_data'].get(field):
                        missing_fields.append(label)
                
                if not missing_fields:
                    session['status'] = 'ready_to_create'
                    return {
                        "status": "success",
                        "message": "Все данные собраны! Документ готов к созданию.",
                        "next_step": "ready_to_create",
                        "user_data": session['user_data']
                    }
                else:
                    # Есть еще недостающие поля, задаем следующие вопросы
                    return self.ask_questions_for_autofill(user_id, selected_template['name'], session['user_data'])
            else:
                return {
                    "status": "success",
                    "message": "Ответ сохранен. Продолжаем сбор данных.",
                    "remaining_questions": len(remaining_questions),
                    "next_step": "answering_questions"
                }
                
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка обработки ответа: {str(e)}"
            }
    
    def finalize_autofill(self, user_id: str, document_name: str) -> Dict:
        """
        Завершает процесс автозаполнения и создает документ
        
        Args:
            user_id: ID пользователя
            document_name: Название документа
        
        Returns:
            Созданный документ
        """
        try:
            # Находим сессию пользователя
            session = self._find_user_session(user_id)
            if not session:
                return {
                    "status": "error",
                    "message": "Сессия автозаполнения не найдена"
                }
            
            selected_template = session['selected_document']
            if not selected_template:
                return {
                    "status": "error",
                    "message": "Документ не выбран"
                }
            
            # Создаем документ
            created_document = self.create_document_from_template(
                user_id=user_id,
                template_id=selected_template['template_id'],
                user_data=session['user_data'],
                conversation_data={
                    "message": f"Интерактивное автозаполнение документа: {document_name}",
                    "response": "Документ создан через интерактивный процесс"
                },
                send_email=True
            )
            
            if created_document and created_document.get("status") == "success":
                # Очищаем сессию
                session['status'] = 'completed'
                
                # Сохраняем в Google Sheets
                try:
                    google_sheets_service.save_document({
                        "user_id": user_id,
                        "full_name": session['user_data'].get("full_name", ""),
                        "email": session['user_data'].get("email", ""),
                        "document_type": self._get_category_from_name(selected_template['name']),
                        "template_name": selected_template['name'],
                        "filepath": created_document.get("filepath", ""),
                        "download_url": created_document.get("download_url", ""),
                        "completeness_score": 100,  # В интерактивном режиме считаем данные полными
                        "confidence_score": 100,
                        "data_quality": "excellent"
                    })
                except Exception as e:
                    print(f"Ошибка сохранения в Google Sheets (не критично): {e}")
                
                return {
                    "status": "success",
                    "message": f"Документ '{selected_template['name']}' успешно создан!",
                    "document": created_document,
                    "user_data": session['user_data']
                }
            else:
                return {
                    "status": "error",
                    "message": "Ошибка создания документа"
                }
                
        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка завершения автозаполнения: {str(e)}"
            }
    
    def _find_user_session(self, user_id: str) -> Optional[Dict]:
        """
        Находит активную сессию автозаполнения для пользователя
        
        Args:
            user_id: ID пользователя
        
        Returns:
            Сессия автозаполнения или None
        """
        print(f"[DEBUG] Поиск сессии для пользователя: {user_id}")
        print(f"[DEBUG] Всего сессий: {len(self.autofill_sessions)}")
        
        for session_id, session in self.autofill_sessions.items():
            print(f"[DEBUG] Проверяем сессию: {session_id}, user_id: {session['user_id']}, status: {session['status']}")
            if session['user_id'] == user_id and session['status'] not in ['completed', 'cancelled']:
                print(f"[DEBUG] Найдена активная сессия: {session_id}")
                return session
        
        print(f"[DEBUG] Активная сессия для пользователя {user_id} не найдена")
        return None
    
    def _generate_question_for_field(self, field: str, label: str, document_name: str) -> str:
        """
        Генерирует вопрос для конкретного поля
        
        Args:
            field: Название поля
            label: Отображаемое название поля
            document_name: Название документа
        
        Returns:
            Текст вопроса
        """
        # Специальные вопросы для конкретных документов
        if 'вступление' in document_name.lower() and field == 'full_name':
            return "Укажите ваше ФИО в родительном падеже (например: Иванова Ивана Ивановича)"
        
        question_templates = {
            "full_name": f"Пожалуйста, укажите ваше полное ФИО в родительном падеже для документа '{document_name}'",
            "email": f"Укажите ваш email адрес для документа '{document_name}'",
            "phone": f"Укажите ваш номер телефона для документа '{document_name}'",
            "organization": f"Укажите название вашей организации для документа '{document_name}'",
            "position": f"Укажите вашу должность для документа '{document_name}'",
            "inn": f"Укажите ваш ИНН для документа '{document_name}'",
            "address": f"Укажите ваш адрес для документа '{document_name}'",
            "passport": f"Укажите ваши паспортные данные для документа '{document_name}'",
            "birth_date": f"Укажите вашу дату рождения для документа '{document_name}'",
            "business_type": f"Укажите тип вашего бизнеса для документа '{document_name}'"
        }
        
        return question_templates.get(field, f"Пожалуйста, укажите {label.lower()} для документа '{document_name}'")
    
    def _get_field_input_type(self, field: str) -> str:
        """
        Определяет тип поля ввода для поля
        
        Args:
            field: Название поля
        
        Returns:
            Тип поля ввода
        """
        input_types = {
            "email": "email",
            "phone": "tel",
            "birth_date": "date",
            "inn": "number"
        }
        
        return input_types.get(field, "text")


assistant_service = AssistantService()

